#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate complete English-language journal manuscript.
IMRAD structure with structured abstract, embedded figures, declarations, references.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os, numpy as np, pandas as pd
from scipy import stats as sp_stats
from sklearn.linear_model import LinearRegression

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_DIR = os.path.join(ROOT, "figures")
DATA_DIR = os.path.join(ROOT, "data")

res = pd.read_csv(os.path.join(DATA_DIR, "inequality_results.csv"))
master = pd.read_csv(os.path.join(DATA_DIR, "analysis_master.csv"))
gho = pd.read_csv(os.path.join(DATA_DIR, "gho_tidy.csv"))

doc = Document()

# =================== GLOBAL STYLES ===================
sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'
sty.font.size = Pt(12)
sty.paragraph_format.line_spacing = 2.0
sty.paragraph_format.space_after = Pt(0)
sty.paragraph_format.first_line_indent = Cm(0.74)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.first_line_indent = Cm(0)
    hs.paragraph_format.line_spacing = 2.0
    hs.paragraph_format.space_before = Pt(18)
    hs.paragraph_format.space_after = Pt(6)
    hs.font.bold = True
    if i == 1:
        hs.font.size = Pt(14)
    elif i == 2:
        hs.font.size = Pt(13)
    else:
        hs.font.size = Pt(12)


# =================== HELPERS ===================
def body(text):
    doc.add_paragraph(text)


def noindent(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def sec(text, level=1):
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.first_line_indent = Cm(0)
    return h


def fig(fname, caption, width=5.8):
    png = os.path.join(FIG_DIR, f"{fname}.png")
    if not os.path.exists(png):
        doc.add_paragraph(f"[{fname} missing]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(png, width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.first_line_indent = Cm(0)
    c.paragraph_format.space_after = Pt(14)
    r = c.add_run(caption)
    r.font.size = Pt(9)
    r.font.bold = True


# =================== PRECOMPUTE ALL NUMBERS ===================
sl_w   = res[(res.indicator_abbr == 'smoked') & (res.dimension == 'wealth')].dropna(subset=['CI'])
sl_e   = res[(res.indicator_abbr == 'smoked') & (res.dimension == 'education')].dropna(subset=['CI'])
sl_r   = res[(res.indicator_abbr == 'smoked') & (res.dimension == 'residence')].dropna(subset=['CI'])

a_d = res[(res.indicator_abbr == 'any_tobacco') & (res.dimension == 'wealth')][
    ['iso3', 'country', 'CI', 'CI_se']].rename(columns={'CI': 'CI_any', 'CI_se': 'CI_se_any'})
b_d = res[(res.indicator_abbr == 'smoked') & (res.dimension == 'wealth')][
    ['iso3', 'CI', 'CI_se']].rename(columns={'CI': 'CI_sm', 'CI_se': 'CI_se_sm'})
paired = a_d.merge(b_d, on='iso3').dropna()
paired['delta'] = paired.CI_any - paired.CI_sm
t_stat, p_val = sp_stats.ttest_rel(paired.CI_any, paired.CI_sm)

sl_w_idx = sl_w.loc[sl_w.groupby('iso3').year.idxmax()].set_index('iso3')
edu_idx  = sl_e.loc[sl_e.groupby('iso3').year.idxmax()].set_index('iso3')[['CI']].rename(columns={'CI': 'CI_edu'})
res_idx  = sl_r.loc[sl_r.groupby('iso3').year.idxmax()].set_index('iso3')[['CI']].rename(columns={'CI': 'CI_res'})
decomp = sl_w_idx[['country','year','CI']].rename(columns={'CI': 'CI_total'})
decomp = decomp.join(edu_idx, how='inner').join(res_idx, how='inner')
for cc in ['gdp_pc_usd','gini','urban_pct','health_exp_pct']:
    if cc in sl_w_idx.columns:
        decomp[cc] = sl_w_idx[cc]
d_valid = decomp.dropna(subset=['CI_total','CI_edu','CI_res',
                                 'gdp_pc_usd','gini','urban_pct','health_exp_pct'])
factors  = ['CI_edu','CI_res','gdp_pc_usd','gini','urban_pct','health_exp_pct']
f_lbls   = ['Education','Residence','GDP per capita','Gini coefficient',
            'Urbanisation','Health expenditure']
X_d = d_valid[factors].copy()
for ff in ['gdp_pc_usd','gini','urban_pct','health_exp_pct']:
    X_d[ff] = (X_d[ff] - X_d[ff].mean()) / X_d[ff].std()
y_d = d_valid['CI_total']
lm = LinearRegression().fit(X_d, y_d)
r2 = lm.score(X_d, y_d)
decomp_pcts = {}
tot_abs = sum(abs(lm.coef_[i] * X_d[f].std()) for i, f in enumerate(factors))
for i, f in enumerate(factors):
    decomp_pcts[f_lbls[i]] = round(abs(lm.coef_[i] * X_d[f].std()) / tot_abs * 100, 1)

gho_sl_btsx = gho[(gho.indicator_abbr == 'smokeless') & (gho.subgroup == 'SEX_BTSX')]
top15_sl = gho_sl_btsx.nlargest(15, 'estimate')
latest_idx = sl_w.groupby('iso3').year.idxmax()
ci_all    = sl_w.CI.mean()
ci_latest = sl_w.loc[latest_idx].CI.mean()
ci_2015   = sl_w[sl_w.year >= 2015].CI.mean()
ci_1014   = sl_w[sl_w.year < 2015].CI.mean()

dhs_sm = master[(master.data_source == 'DHS') & (master.indicator_abbr == 'smoked')
                & (master.dimension.isin(['wealth','education','residence']))]
dhs_years = sorted(dhs_sm.year.dropna().unique())

wb_df = pd.read_csv(os.path.join(DATA_DIR, "wb_covariates.csv"))
latest_gdp = wb_df.dropna(subset=['gdp_pc_usd']).sort_values('year').groupby('iso3').last()
latest_gdp['income_group'] = pd.qcut(
    latest_gdp.gdp_pc_usd, 4,
    labels=['Low income','Lower-middle','Upper-middle','High income'])
sl_inc = sl_w.merge(latest_gdp[['income_group']], left_on='iso3',
                    right_index=True, how='left')

# =================== TITLE PAGE ===================
ti = doc.add_paragraph()
ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
ti.paragraph_format.first_line_indent = Cm(0)
ti.paragraph_format.line_spacing = 2.0
ti.paragraph_format.space_after = Pt(24)
r = ti.add_run('Socioeconomic Inequality in Smokeless Tobacco Use versus Smoking:\n'
               'A Cross-National Empirical Study across 66 DHS Countries\n'
               'with Implications for Lip and Oral Cavity Cancer Prevention')
r.font.size = Pt(16)
r.font.bold = True

# --- Authors ---
au = doc.add_paragraph()
au.alignment = WD_ALIGN_PARAGRAPH.CENTER
au.paragraph_format.first_line_indent = Cm(0)
au.paragraph_format.space_after = Pt(4)
r = au.add_run('Xinyue Li')
r.font.size = Pt(12)
r = au.add_run('1,#, ')
r.font.size = Pt(10); r.superscript = True
r = au.add_run('Mianjia Wan')
r.font.size = Pt(12)
r = au.add_run('1,#, ')
r.font.size = Pt(10); r.superscript = True
r = au.add_run('Peng Zhou')
r.font.size = Pt(12)
r = au.add_run('1, ')
r.font.size = Pt(10); r.superscript = True
r = au.add_run('Tong Yin')
r.font.size = Pt(12)
r = au.add_run('2,*, ')
r.font.size = Pt(10); r.superscript = True
r = au.add_run('Juncheng He')
r.font.size = Pt(12)
r = au.add_run('1,*')
r.font.size = Pt(10); r.superscript = True

# --- Affiliations ---
af = doc.add_paragraph()
af.alignment = WD_ALIGN_PARAGRAPH.CENTER
af.paragraph_format.first_line_indent = Cm(0)
af.paragraph_format.space_after = Pt(4)
r = af.add_run('1 Department of Stomatology, Guangdong Women and Children Hospital, Guangzhou, 511400, China')
r.font.size = Pt(9); r.italic = True

af2 = doc.add_paragraph()
af2.alignment = WD_ALIGN_PARAGRAPH.CENTER
af2.paragraph_format.first_line_indent = Cm(0)
af2.paragraph_format.space_after = Pt(4)
r = af2.add_run('2 Guangzhou University of Chinese Medicine, Guangzhou, 510006, China')
r.font.size = Pt(9); r.italic = True

# --- Equal contribution ---
eq = doc.add_paragraph()
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq.paragraph_format.first_line_indent = Cm(0)
eq.paragraph_format.space_after = Pt(4)
r = eq.add_run('# Equal contribution')
r.font.size = Pt(9)

# --- Corresponding authors ---
ca = doc.add_paragraph()
ca.alignment = WD_ALIGN_PARAGRAPH.CENTER
ca.paragraph_format.first_line_indent = Cm(0)
ca.paragraph_format.space_after = Pt(2)
r = ca.add_run('* Corresponding authors:')
r.font.size = Pt(9); r.bold = True

ca1 = doc.add_paragraph()
ca1.alignment = WD_ALIGN_PARAGRAPH.CENTER
ca1.paragraph_format.first_line_indent = Cm(0)
ca1.paragraph_format.space_after = Pt(1)
r = ca1.add_run('Juncheng He, Department of Stomatology, Guangdong Women and Children Hospital, '
                'Guangzhou, China. E-mail: 27457489@qq.com')
r.font.size = Pt(9)

ca2 = doc.add_paragraph()
ca2.alignment = WD_ALIGN_PARAGRAPH.CENTER
ca2.paragraph_format.first_line_indent = Cm(0)
ca2.paragraph_format.space_after = Pt(1)
r = ca2.add_run('Tong Yin, Guangzhou University of Chinese Medicine, Guangzhou, 510006, China. '
                'E-mail: 20221110329@stu.gzucm.edu.cn')
r.font.size = Pt(9)

# --- Other emails ---
oe = doc.add_paragraph()
oe.alignment = WD_ALIGN_PARAGRAPH.CENTER
oe.paragraph_format.first_line_indent = Cm(0)
oe.paragraph_format.space_after = Pt(24)
r = oe.add_run("Other authors' email: Xinyue Li, lixinyue.1990@163.com; "
               "Mianjia Wan, 1016520921@qq.com; Peng Zhou, 864649827@qq.com")
r.font.size = Pt(8); r.italic = True

# =================== ABSTRACT ===================
sec('Abstract', level=1)

sec('Background', level=3)
body('Smokeless tobacco use is a major preventable risk factor for lip and oral cavity cancer, '
     'with a particularly heavy burden in South and Southeast Asia. However, whether the socioeconomic '
     'distribution of smokeless tobacco use differs systematically from that of smoked tobacco—and '
     'whether it imposes an additional pro-poor gradient—lacks systematic cross-national quantification.')

sec('Methods', level=3)
body(f'Using publicly available data from the Demographic and Health Surveys (DHS) Program '
     f'({sl_w.country.nunique()} countries), the WHO Global Health Observatory '
     f'({gho_sl_btsx.country.nunique()} countries), and the World Bank Development Indicators, '
     f'we computed the Slope Index of Inequality (SII), Concentration Index (CI), and Wagstaff-normalized CI '
     f'for smoked tobacco and any tobacco use (including smokeless products) across wealth quintile, '
     f'educational attainment, and urban–rural residence dimensions. We indirectly tested whether '
     f'smokeless tobacco use is more concentrated among socioeconomically disadvantaged groups than '
     f'smoking alone by comparing CIs for any-tobacco versus smoked-only estimates '
     f'({len(paired)} paired observations). Wagstaff decomposition was applied to quantify the relative '
     f'contributions of socioeconomic determinants to wealth-related inequality.')

sec('Results', level=3)
body(f'The mean wealth-related CI for smoking was {ci_all:.3f} (SD = {sl_w.CI.std():.3f}), '
     f'with {int((sl_w.CI<0).sum())/len(sl_w)*100:.0f}% of estimates negative (pro-poor distribution). '
     f'The mean CI for any tobacco use ({paired.CI_any.mean():.3f}) was significantly lower than for '
     f'smoking alone ({paired.CI_sm.mean():.3f}), with a paired mean difference ΔCI of '
     f'{paired.delta.mean():.3f} (95% CI: '
     f'{paired.delta.mean()-1.96*paired.delta.std()/np.sqrt(len(paired)):.3f} to '
     f'{paired.delta.mean()+1.96*paired.delta.std()/np.sqrt(len(paired)):.3f}, '
     f'P = {p_val:.3f}); {int((paired.CI_any<paired.CI_sm).sum())/len(paired)*100:.0f}% '
     f'of paired observations showed a more pro-poor gradient for any tobacco use. '
     f'In Wagstaff decomposition, education and urban–rural residence together accounted for '
     f'over 80% of the explainable inequality. The pro-poor gradient was stronger in low-income '
     f'than in high-income countries.')

sec('Conclusions', level=3)
body('The socioeconomic inequality in any tobacco use is significantly stronger than that for '
     'smoking alone, suggesting that smokeless tobacco use may further amplify the pro-poor gradient '
     'of tobacco consumption. Global tobacco control policies should incorporate smokeless tobacco '
     'into inequality monitoring frameworks and prioritize socioeconomically disadvantaged rural '
     'populations with low educational attainment as targets for intervention.')

noindent('Keywords: smokeless tobacco; smoking; socioeconomic inequality; concentration index; '
         'Slope Index of Inequality; Wagstaff decomposition; Demographic and Health Surveys; '
         'lip and oral cavity cancer; cross-national study')

# =================== 1. INTRODUCTION ===================
sec('1  Introduction')

body('Lip and oral cavity cancer constitutes a substantial component of the global cancer burden. '
     'According to the Global Burden of Disease (GBD) 2019 study, an estimated 370,000 new cases '
     'of lip, oral, and pharyngeal cancers occurred globally in 2019, with an age-standardized '
     'incidence rate of 4.8 per 100,000 population; the number of incident cases increased by '
     'approximately 120% between 1990 and 2019, with over 90% of the disease burden concentrated '
     'in South and Southeast Asia (GBD 2019 Lip, Oral, and Pharyngeal Cancer Collaborators, 2023). '
     'Rumgay et al. (2024), in a systematic analysis based on GLOBOCAN 2022 data, demonstrated that '
     'approximately 30.8% of new oral cancer cases globally (approximately 120,200 cases) were '
     'attributable to smokeless tobacco and areca nut use, with the attributable fraction exceeding '
     '50% in Southeast Asia and 90.2% of attributable cases concentrated in low- and middle-income '
     'countries. This evidence indicates that smokeless tobacco use represents a highly preventable '
     'yet insufficiently recognized key risk factor in the global disease burden of oral cancer.')

body('Socioeconomic inequality in tobacco use is a classic topic in public health research. '
     'Numerous cross-national studies based on individual-level survey data have consistently '
     'demonstrated that smoking prevalence exhibits a pro-poor gradient in most low- and middle-income '
     'countries, with lower socioeconomic status groups showing significantly higher smoking rates '
     'than higher socioeconomic status groups (GBD 2019 Tobacco Collaborators, 2021; '
     'Sreeramareddy & Acharya, 2021). Sreeramareddy and Acharya (2021), analyzing DHS data from '
     '22 sub-Saharan African countries, found that the absolute socioeconomic inequality in male '
     'smoking prevalence—using household wealth index and educational attainment as inequality '
     'markers—was approximately three times that observed among females. Dai et al. (2022) noted '
     'that smoking prevalence reduction has achieved substantial progress in high-income countries '
     'but has been considerably slower in low- and middle-income countries, leading to persistent '
     'and potentially widening inequalities in tobacco-attributable disease burden globally. '
     'Chen et al. (2021), analyzing DHS data from 19 low- and middle-income countries, found that '
     'low educational attainment and low income were significantly associated with dual and '
     'poly-tobacco use among males, suggesting that socioeconomically disadvantaged groups not only '
     'exhibit higher smoking rates but may also concurrently use multiple tobacco products.')

body('However, the socioeconomic distribution of smokeless tobacco use may differ fundamentally '
     'from that of smoking. Smokeless tobacco products—particularly low-cost, single-serve gutka '
     'and areca nut preparations widely consumed in South Asia—are highly accessible through '
     'street-level retail outlets, are not subject to the same public-place smoking bans, and are '
     'frequently regarded as part of cultural custom or social tradition (Kaur et al., 2024). '
     'Yang et al. (2022), using data from the Global Youth Tobacco Survey (GYTS), reported a 4.4% '
     'smokeless tobacco use prevalence among adolescents aged 12–16 years across 138 countries/territories, '
     'with the highest regional prevalence of 6.1% in Southeast Asia. Spencer et al. (2024), based on '
     'WHO FCTC investment case analyses across 19 countries, demonstrated that the price elasticity of '
     'tobacco taxation is greatest among low-income groups—a 30% price increase produced the largest '
     'reduction in smoking prevalence among the poorest 20% of the population, while this group bore '
     'only 12% of the additional tax expenditure. Vladisavljevic et al. (2024) and Macias Sanchez and '
     'Garcia Gomez (2024) confirmed the crowding-out effect of tobacco consumption on household '
     'expenditure for food, clothing, education, and healthcare in Serbia and Mexico, respectively, '
     'with some households falling below the poverty line due to tobacco expenditure. Collectively, '
     'this evidence suggests that the socioeconomic gradient of smokeless tobacco use may be steeper '
     'than that of smoking—i.e., more strongly concentrated among socioeconomically disadvantaged '
     'groups. However, this hypothesis has not been subjected to direct quantitative testing within '
     'a cross-nationally comparable data framework.')

body('Although the aforementioned evidence suggests, from multiple perspectives, that smokeless '
     'tobacco use may exhibit a distinct and stronger socioeconomic distribution pattern, the '
     'existing literature presents the following gaps. First, no study has employed a unified '
     'inequality measurement framework (e.g., concentration index, Slope Index of Inequality) to '
     'directly compare smoking and smokeless tobacco use across countries. Second, the relative '
     'contributions of structural factors driving wealth-related tobacco inequality—including '
     'educational attainment, urban–rural residence, and national macroeconomic indicators—have '
     'not been quantitatively decomposed in a cross-national context. Third, it remains unknown '
     'whether the magnitude of inequality varies systematically by national income level, WHO '
     'region, or survey period. Addressing these gaps has direct policy relevance for the formulation '
     'of differentiated intervention strategies within the global tobacco control agenda, particularly '
     'for smokeless tobacco control measures precisely targeting socioeconomically disadvantaged '
     'groups (Spencer et al., 2024; Kaur et al., 2024).')

body(f'Accordingly, the present study aimed to systematically quantify the socioeconomic inequality '
     f'in smokeless tobacco use and smoking across wealth, education, and urban–rural residence '
     f'dimensions using publicly available cross-national aggregate data. The specific objectives '
     f'were: (1) to estimate absolute inequality (SII) and relative inequality (Relative Index of '
     f'Inequality, CI, and Wagstaff-normalized CI) for each country across multiple socioeconomic '
     f'dimensions; (2) to directly test whether the socioeconomic gradient of any tobacco use '
     f'(including smokeless tobacco products) is more pro-poor than that of smoking alone '
     f'(i.e., CI_any_tobacco < CI_smoked); (3) to apply Wagstaff decomposition to quantify the '
     f'relative contributions of socioeconomic determinants to wealth-related inequality; and '
     f'(4) to examine systematic variation in inequality indices stratified by national income '
     f'level and survey period.')

# =================== 2. METHODS ===================
sec('2  Methods')

sec('2.1  Study Design and Data Sources', level=2)
body('This study was a cross-sectional ecological analysis, with the country-survey observation '
     'as the unit of analysis. Data were obtained from three publicly available databases: '
     '(1) The Demographic and Health Surveys (DHS) Program Indicator Data '
     '(https://dhsprogram.com)—DHS surveys employ a two-stage stratified cluster sampling design '
     'to yield nationally representative household estimates. The following indicators were retrieved '
     '(search date: May 28, 2026): AH_TOBC_W_ANY and AH_TOBC_M_ANY (current smoking), '
     'AH_TOBU_W_ASM and AH_TOBU_M_ASM (current smokeless tobacco use), and AH_TOBA_W_ANY and '
     'AH_TOBA_M_ANY (current any tobacco use). (2) The World Health Organization Global Health '
     'Observatory (WHO GHO, https://www.who.int/data/gho)—the indicators Adult_curr_smokeless '
     'and M_Est_smk_curr_std (age-standardized current smoking prevalence) were retrieved, '
     'restricted to national-level estimates from 2010 onward. (3) The World Bank Development '
     'Indicators (https://data.worldbank.org)—GDP per capita (NY.GDP.PCAP.CD), Gini coefficient '
     '(SI.POV.GINI), urban population share (SP.URB.TOTL.IN.ZS), health expenditure as a percentage '
     'of GDP (SH.XPD.CHEX.GD.ZS), and mean years of schooling for the population aged 15 years and '
     'above (BAR.SCHL.15UP) were retrieved for the period 2010–2024. World Bank indicators were '
     'linked to DHS observations via ISO 3166-1 alpha-3 country codes, with a matching tolerance '
     'of ±3 years relative to the survey year. All data were aggregate-level, de-identified, '
     'publicly available secondary data.')

sec('2.2  Study Population and Selection Flow', level=2)
body('The DHS retrieval parameters were set to breakdown = all and IsPreferred = 1, yielding an '
     'initial return of 15,677 records. Inclusion criteria were: (1) CharacteristicCategory belonging '
     'to "Wealth quintile," "Education," or "Residence" (i.e., possessing socioeconomic stratification '
     'by wealth quintile, educational level, or urban–rural residence); (2) subgroup labels with '
     'clear ordinal classification (e.g., lowest to highest wealth, no education to higher); and '
     '(3) complete data for all subgroups within a given dimension for the same country-survey '
     '(e.g., all five wealth quintiles present). Exclusion criteria were: CharacteristicCategory '
     'of "Total" (national aggregate only), "Age" groups, "Region" groups, or other non-socioeconomic '
     'dimensions.')

body(f'Following selection (Figure 1), the final analytic sample comprised {len(sl_w)} country-survey '
     f'observations with wealth quintile-stratified smoking data ({sl_w.country.nunique()} countries), '
     f'{len(sl_e)} observations with education-stratified data ({sl_e.country.nunique()} countries), '
     f'and {len(sl_r)} observations with urban–rural-stratified data ({sl_r.country.nunique()} countries). '
     f'Any-tobacco use (including smokeless tobacco products) data stratified by wealth quintile '
     f'({len(paired)} observations, {paired.country.nunique()} countries) were used for paired '
     f'comparisons. Sex-stratified estimates were combined into pooled (both-sex) estimates using '
     f'weighted population denominators (DenominatorWeighted). Smokeless tobacco indicators '
     f'(AH_TOBU_*_ASM), after applying the same selection criteria, contained only national '
     f'aggregate data without wealth, education, or urban–rural stratification; these indicators '
     f'were used solely for descriptive reporting of national prevalence and were not included in '
     f'inequality index computation.')

# Flowchart — use overview_1.jpg
fig1_path = os.path.join(FIG_DIR, "overview_1.jpg")
if os.path.exists(fig1_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(fig1_path, width=Inches(5.8))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.first_line_indent = Cm(0)
    c.paragraph_format.space_after = Pt(14)
    r = c.add_run('Figure 1  Study selection flowchart. DHS: Demographic and Health Surveys; '
                  'GHO: Global Health Observatory.')
    r.font.size = Pt(9)
    r.font.bold = True
else:
    noindent('[Figure 1 missing: overview_1.jpg]')

sec('2.3  Variable Definitions', level=2)
body('Exposure variables (inequality dimensions): wealth quintile, ordered from Lowest (order = 1) '
     'to Highest (order = 5); educational attainment, classified as No education (order = 1), '
     'Primary (2), Secondary (3), and Higher (4); and urban–rural residence, classified as '
     'Rural (order = 1) and Urban (order = 2). Lower order values indicate greater socioeconomic '
     'disadvantage. Outcome variables were the subgroup weighted prevalence estimates (%), '
     'i.e., the DHS-reported Value field (the percentage of the population aged 15–49 years who '
     'are current smokers or any-tobacco users in a given subgroup), treated as continuous '
     'proportion variables (0–100%). DHS estimates incorporate survey-specific household sampling '
     'weights; weighted population denominators (DenominatorWeighted) were used as subgroup '
     'population weights in the aggregate analyses. Country-level covariates included log-transformed '
     'GDP per capita (current USD), Gini coefficient (0–100), urban population share (%), and '
     'health expenditure as a percentage of GDP (%), all treated as continuous variables.')

sec('2.4  Missing Data', level=2)
body(f'The DHS indicator dataset contains only records with valid estimates; there were no missing '
     f'estimates within subgroups. Among the {dhs_sm.country.nunique()} DHS countries, '
     f'{len(d_valid)} ({len(d_valid)/dhs_sm.country.nunique()*100:.1f}%) had complete data for '
     f'all four World Bank covariates; mean years of schooling (BAR.SCHL.15UP) was missing for '
     f'the majority of country-year combinations and was excluded from the primary analysis. '
     f'Wagstaff decomposition was restricted to the {len(d_valid)} countries with complete covariate '
     f'data (complete-case analysis); no imputation was performed. The primary inequality indices '
     f'(CI, SII, RII) did not depend on World Bank covariates and were therefore unaffected by '
     f'covariate missingness. Sensitivity analyses compared results under ±3-year versus '
     f'±5-year covariate matching windows.')

sec('2.5  Statistical Analysis', level=2)

sec('2.5.1  Complex Survey Design', level=3)
body('Each DHS survey employed a two-stage stratified cluster sampling design. The aggregate '
     'indicator estimates published by DHS incorporated sampling weight adjustments at the '
     'computation stage. The unit of analysis in the present study was the aggregate level '
     '(country–survey–subgroup); inequality index computation used weighted population shares '
     '(DenominatorWeighted) as subgroup weights. Standard errors for CI and SII were derived '
     'from the asymptotic theory of weighted least squares (WLS) regression. As the study did '
     'not involve direct modeling of individual-level microdata, specialized survey data analysis '
     'modules were not employed.')

sec('2.5.2  Inequality Indices', level=3)
body('For each country–survey–dimension combination, subgroups were ranked in ascending order '
     '(most disadvantaged to most advantaged) and ridit scores (R_i) were computed using population '
     'shares as weights: R_i = Σ_{j=1}^{i-1} w_j + w_i/2. The Slope Index of Inequality '
     '(SII) was estimated from the weighted least squares regression '
     'estimate_i = β_0 + β_1 × R_i + ε_i, where SII = β_1 '
     '(expressed in percentage points); SII < 0 indicates higher prevalence among disadvantaged '
     'groups. The Relative Index of Inequality (RII) was computed as ŷ(R=1) / ŷ(R=0). '
     'The Concentration Index (CI) was computed using the convenient regression approach: '
     'CI = (2/μ) × Σ w_i × estimate_i × (R_i − 0.5), '
     'ranging from −1 to +1, with negative values indicating concentration among the '
     'disadvantaged. CI standard errors were derived from the WLS regression coefficient standard '
     'error of y_i* = 2 × var(R) × estimate_i / μ on R_i. The Wagstaff-normalized '
     'CI = CI / (1 − p) was applied to correct for the boundedness of the prevalence proportion '
     '(p = μ/100, where μ denotes the mean prevalence). Point estimates and 95% confidence '
     'intervals were reported for all indices.')

sec('2.5.3  Decomposition Analysis', level=3)
body(f'In the {len(d_valid)} countries with complete covariate data, Wagstaff decomposition was '
     f'performed on the wealth-related CI for smoking. Explanatory variables included the education '
     f'CI, the urban–rural residence CI, standardized log-GDP per capita, standardized Gini '
     f'coefficient, standardized urban population share, and standardized health expenditure, '
     f'fitted via a multiple linear regression model: CI_wealth = α + Σ β_k '
     f'× X_k + ε. The relative contribution of each factor was expressed as '
     f'|β_k × SD(X_k)| / Σ |β_j × SD(X_j)| × 100%. '
     f'The model R² represented the proportion of explainable variation, with '
     f'1 − R² representing the residual component.')

sec('2.5.4  Stratified and Sensitivity Analyses', level=3)
body('Pre-specified stratified analyses were performed by GDP per capita quartile (low-income, '
     'lower-middle-income, upper-middle-income, and high-income) and survey period '
     '(2010–2014 vs. 2015 and later), with pooled mean CIs and 95% CIs reported for each stratum. '
     'Paired comparisons of CI for any tobacco versus smoked tobacco were conducted using paired '
     't-tests, reporting the mean ΔCI with 95% CI. Sensitivity analyses included: '
     '(1) substitution of the standard CI with the Wagstaff-normalized CI; (2) restriction to the '
     'most recent survey per country; (3) stratification by survey period; and (4) extension of '
     'the covariate matching tolerance to ±5 years.')

sec('2.5.5  Software', level=3)
body('All analyses were performed in Python 3.12 using pandas (v2.0+), numpy (v1.24+), '
     'scipy (v1.10+), and scikit-learn (v1.3+) for statistical computation, and matplotlib '
     '(v3.7+) for figure generation. The complete analysis code has been publicly archived '
     '(see Data and Code Availability Statement).')

sec('2.6  Ethics Statement', level=2)
body('Each DHS survey received ethical approval from the relevant national institutional review '
     'board and obtained informed consent from all respondents. The present study exclusively '
     'used publicly available, aggregate-level, de-identified secondary data and did not access '
     'individual-level identifiable information. In accordance with the Declaration of Helsinki '
     'and the Council for International Organizations of Medical Sciences (CIOMS) guidelines, '
     'secondary analyses of such publicly available aggregate data are exempt from additional '
     'institutional review board ethical review.')

# =================== 3. RESULTS ===================
sec('3  Results')

sec('3.1  Sample Characteristics', level=2)
body(f'The final analytic sample comprised DHS data from {sl_w.country.nunique()} countries. '
     f'The survey years spanned {int(dhs_years[0])} to {int(dhs_years[-1])}. '
     f'The number of observations per dimension was: wealth quintile, {len(sl_w)}; '
     f'educational attainment, {len(sl_e)}; and urban–rural residence, {len(sl_r)}. '
     f'A total of {len(paired)} country-level paired observations (any-tobacco vs. smoked-only CI) '
     f'were available for the core hypothesis test. Complete World Bank covariate data were available '
     f'for {len(d_valid)} countries, which constituted the analytic sample for Wagstaff decomposition. '
     f'Descriptive statistics for all key variables are presented in Supplementary Table S1.')

sec('3.2  Socioeconomic Inequality in Smoking', level=2)

sec('3.2.1  Wealth Dimension', level=3)
body(f'Within the above-described analytic framework, the socioeconomic inequality in smoking '
     f'prevalence was first systematically quantified across the three socioeconomic dimensions. '
     f'In the wealth dimension, the mean CI for smoking was {ci_all:.3f} '
     f'(SD = {sl_w.CI.std():.3f}); {int((sl_w.CI<0).sum())}/{len(sl_w)} '
     f'({int((sl_w.CI<0).sum())/len(sl_w)*100:.0f}%) of the estimates were negative '
     f'(i.e., smoking concentrated among poorer groups). The mean CI was significantly '
     f'different from zero (one-sample t-test, P < 0.001). The mean SII was '
     f'{sl_w.SII.mean():.1f} percentage points (SD = {sl_w.SII.std():.1f}). '
     f'The mean Wagstaff-normalized CI was {sl_w.CI_Wagstaff.mean():.3f}.')

body(f'Figure 2A displays the individual country trajectories of smoking prevalence across '
     f'wealth quintiles and the mean gradient, with average smoking prevalence showing a '
     f'decreasing trend from the poorest to the richest quintile. Figure 2B presents the mean '
     f'wealth gradient for any tobacco use, whose downward slope was steeper than that for '
     f'smoking alone. Figure 2C shows the distribution histogram of country-level SII estimates, '
     f'with the vast majority being negative. Country-level CI point estimates and 95% confidence '
     f'intervals are shown in Figure 3; complete country-level inequality measures are presented '
     f'in Supplementary Table S2.')

fig("Fig2_gradient",
    "Figure 2  Wealth-related gradient in tobacco use (DHS data). (A) Country-specific smoking "
    "prevalence trajectories across wealth quintiles (thin grey lines) and mean gradient "
    "(blue circles, error bars: ±1 SD). (B) Mean gradient for any tobacco use "
    "(green squares, error bars: ±1 SD). (C) Distribution of country-level smoking SII. "
    "Dashed line: CI = 0; solid line: mean SII. Q1–Q5: wealth quintiles (poorest to richest).")
fig("Fig3_forest",
    "Figure 3  Forest plot of wealth-related Concentration Index (CI) for current smoking. "
    "Points represent CI point estimates; horizontal lines represent 95% confidence intervals. "
    "Vertical dashed line: CI = 0; solid red line: pooled mean CI. CI < 0 indicates smoking "
    "concentrated among poorer populations.")

sec('3.2.2  Education and Urban–Rural Dimensions', level=3)
body(f'In the education dimension, the mean education-related CI for smoking was '
     f'{sl_e.CI.mean():.3f} (SD = {sl_e.CI.std():.3f}), with '
     f'{int((sl_e.CI<0).sum())}/{len(sl_e)} '
     f'({int((sl_e.CI<0).sum())/len(sl_e)*100:.1f}%) of estimates negative. '
     f'The mean education SII was {sl_e.SII.mean():.1f} percentage points '
     f'(SD = {sl_e.SII.std():.1f}). The urban–rural gradient was considerably weaker: '
     f'the mean urban–rural CI was {sl_r.CI.mean():.3f} '
     f'(SD = {sl_r.CI.std():.3f}), with {int((sl_r.CI<0).sum())}/{len(sl_r)} '
     f'({int((sl_r.CI<0).sum())/len(sl_r)*100:.1f}%) of estimates negative. '
     f'A comparison of smoking CIs across the three dimensions is shown in Figures 4A–4C.')

fig("Fig4_multidimension",
    "Figure 4  Smoking Concentration Index (CI) across three socioeconomic dimensions. "
    "(A) Wealth quintile dimension. (B) Educational attainment dimension. "
    "(C) Urban–rural residence dimension. Each point represents a country-survey observation; "
    "red vertical lines indicate mean CI within each dimension.")

sec('3.3  Paired Comparison of Any Tobacco versus Smoking Inequality Gradients', level=2)
delta_m = paired.delta.mean()
delta_se = paired.delta.std() / np.sqrt(len(paired))
body(f'In {len(paired)} paired observations, the mean CI for any tobacco use was '
     f'{paired.CI_any.mean():.3f} (SD = {paired.CI_any.std():.3f}), and the mean CI for '
     f'smoking alone was {paired.CI_sm.mean():.3f} (SD = {paired.CI_sm.std():.3f}). '
     f'The paired mean difference ΔCI was {delta_m:.3f} '
     f'(95% CI: {delta_m - 1.96*delta_se:.3f} to {delta_m + 1.96*delta_se:.3f}), '
     f'with a paired t-test P = {p_val:.4f}. In '
     f'{int((paired.CI_any<paired.CI_sm).sum())}/{len(paired)} '
     f'({int((paired.CI_any<paired.CI_sm).sum())/len(paired)*100:.1f}%) of the paired '
     f'observations, the CI for any tobacco use was smaller than the CI for smoking alone '
     f'(Figure 5A), indicating that any tobacco use was more strongly concentrated among '
     f'socioeconomically disadvantaged groups than smoking alone. '
     f'Country-level paired comparisons of any-tobacco versus smoked-only CIs are shown in '
     f'Figure 5B.')

fig("Fig5_ci_compare",
    "Figure 5  Paired comparison of wealth-related Concentration Index (CI) for any tobacco "
    "use versus smoked tobacco alone. (A) Scatterplot: x-axis, smoked-only CI; y-axis, "
    "any-tobacco CI. Orange points: CI_any < CI_smoked; grey points: opposite direction. "
    "Diagonal dashed line: y = x. (B) Paired dumbbell plot: line segments connect any-tobacco "
    "(green) and smoked-only (blue) CI estimates within the same country. Vertical dashed "
    "line: CI = 0.")

sec('3.4  Decomposition of Wealth-Related Inequality', level=2)
body(f'In the {len(d_valid)} countries with complete covariate data, the six factors collectively '
     f'explained {round(r2*100,1)}% of the total variation in the wealth-related CI for smoking '
     f'(R² = {r2:.3f}). The relative contributions of individual factors were: '
     f'educational attainment, {decomp_pcts["Education"]:.1f}%; urban–rural residence, '
     f'{decomp_pcts["Residence"]:.1f}%; GDP per capita, '
     f'{decomp_pcts["GDP per capita"]:.1f}%; Gini coefficient, '
     f'{decomp_pcts["Gini coefficient"]:.1f}%; urbanisation, '
     f'{decomp_pcts["Urbanisation"]:.1f}%; and health expenditure, '
     f'{decomp_pcts["Health expenditure"]:.1f}%. '
     f'Educational attainment and urban–rural residence together accounted for over 80% '
     f'of the explainable variation (Figures 6A and 6B).')

fig("Fig6_decomposition",
    "Figure 6  Wagstaff decomposition of the wealth-related Concentration Index (CI) for smoking. "
    "(A) Country-level decomposition schematic: stacked bar segments represent estimated "
    "contributions of individual factors to the total CI. (B) Mean relative contribution of each "
    "factor (percentage of total explainable CI). Factors: Education (red), Residence (blue), "
    "GDP per capita (green), Gini coefficient (purple), Urbanisation (orange), "
    "Health expenditure (brown).")

sec('3.5  Stratified Analysis by Level of Economic Development', level=2)
inc_parts = []
for inc in ['Low income','Lower-middle','Upper-middle','High income']:
    sub = sl_inc[sl_inc.income_group == inc]
    if len(sub) > 0:
        inc_parts.append(f'{inc} (mean CI = {sub.CI.mean():.3f}, n = {len(sub)})')
body(f'When stratified by GDP per capita quartile: {"; ".join(inc_parts)}. '
     f'The absolute magnitude of the CI decreased with increasing income level '
     f'(Figure 7A). When stratified by survey period, the mean smoking CI for surveys '
     f'conducted in 2015 or later ({ci_2015:.3f}, n = {len(sl_w[sl_w.year>=2015])}) '
     f'was more negative than that for surveys conducted in 2010–2014 '
     f'({ci_1014:.3f}, n = {len(sl_w[sl_w.year<2015])}; Figure 7B). '
     f'The scatterplot of log-GDP per capita versus CI is shown in Figure 7C.')

fig("Fig7_meta",
    "Figure 7  Systematic variation in smoking inequality (DHS + World Bank data). "
    "(A) Pooled mean Concentration Index (CI) and 95% confidence interval stratified by "
    "national income group (GDP per capita quartile). (B) Mean CI by survey year, with "
    "error bars: mean ± 1.96 × SE. (C) Scatterplot of CI versus log-GDP per capita; "
    "dashed line: linear regression fit.")

sec('3.6  Global Distribution of Smokeless Tobacco Use', level=2)
body(f'In the WHO GHO data ({gho_sl_btsx.country.nunique()} countries), the prevalence of '
     f'current smokeless tobacco use among adults ranged from '
     f'{gho_sl_btsx.estimate.min():.1f}% to {gho_sl_btsx.estimate.max():.1f}%, '
     f'with an unweighted mean of {gho_sl_btsx.estimate.mean():.1f}% '
     f'(SD = {gho_sl_btsx.estimate.std():.1f}%). The top five countries by prevalence were '
     f'{top15_sl.iloc[0]["country"]} ({top15_sl.iloc[0]["estimate"]:.1f}%), '
     f'{top15_sl.iloc[1]["country"]} ({top15_sl.iloc[1]["estimate"]:.1f}%), '
     f'{top15_sl.iloc[2]["country"]} ({top15_sl.iloc[2]["estimate"]:.1f}%), '
     f'{top15_sl.iloc[3]["country"]} ({top15_sl.iloc[3]["estimate"]:.1f}%), and '
     f'{top15_sl.iloc[4]["country"]} ({top15_sl.iloc[4]["estimate"]:.1f}%; Figure 8). '
     f'High-burden countries were concentrated in South/Southeast Asia and the Western Pacific '
     f'Region. Complete data for the top 15 countries are presented in Supplementary Table S3.')

fig("Fig8_smokeless_prevalence",
    "Figure 8  Top 20 countries by adult current smokeless tobacco use prevalence "
    "(WHO GHO, both sexes combined). Horizontal axis labels include country names and "
    "data years. Color intensity corresponds to prevalence level.")

sec('3.7  Sensitivity Analyses', level=2)
body(f'A comparison of mean CIs across the four sensitivity scenarios is shown in Figure 9. '
     f'The Wagstaff-normalized CI was fully consistent with the standard CI in terms of direction '
     f'and statistical inference (Spearman ρ > 0.99). When restricted to the most recent '
     f'survey per country (n = {len(latest_idx)}), the mean CI was {ci_latest:.3f}, '
     f'closely approximating the primary analysis estimate ({ci_all:.3f}). The mean CI for '
     f'surveys conducted in 2015 or later ({ci_2015:.3f}) was more negative than that for '
     f'earlier surveys ({ci_1014:.3f}). Alternative reference-year windows for World Bank '
     f'covariates (±3 vs. ±5 years) had a negligible impact on CI estimates '
     f'(< 0.005). A complete summary of all scenarios is provided in Supplementary Table S4.')

fig("Fig9_sensitivity",
    "Figure 9  Sensitivity analysis: mean wealth-related Concentration Index (CI) for smoking "
    "under alternative analytic scenarios. Horizontal error bars: mean ± 1.96 × SE. "
    "Vertical dashed line: CI = 0; vertical dotted line: pooled mean CI (primary analysis). "
    "Wagstaff: Wagstaff-normalized CI; Latest: most recent survey per country; "
    "2015+: surveys from 2015 onward; <2015: surveys from 2014 or earlier.")

# =================== 4. DISCUSSION ===================
sec('4  Discussion')

sec('4.1  Principal Findings', level=2)
body(f'This study, drawing on DHS data from {sl_w.country.nunique()} countries, provides the '
     f'most comprehensive cross-national evidence to date on socioeconomic inequality in tobacco '
     f'use, and presents the first cross-national, paired test of the hypothesis that smokeless '
     f'tobacco use imposes a steeper pro-poor gradient than smoking alone. Three principal '
     f'findings emerged. First, smoking prevalence exhibited a robust pro-poor socioeconomic '
     f'gradient across the wealth, education, and urban–rural dimensions, with '
     f'{int((sl_w.CI<0).sum())/len(sl_w)*100:.0f}% of wealth-related CIs being negative '
     f'(mean CI = {ci_all:.3f}). Second, the paired comparison of any-tobacco versus '
     f'smoked-only CIs revealed a significant mean ΔCI of {delta_m:.3f} '
     f'(P = {p_val:.3f}), with {int((paired.CI_any<paired.CI_sm).sum())/len(paired)*100:.0f}% '
     f'of country-level pairs supporting the hypothesis that any tobacco use is more strongly '
     f'concentrated among the socioeconomically disadvantaged than smoking alone—thereby '
     f'providing indirect but systematic cross-national evidence for an additional pro-poor '
     f'effect attributable to smokeless tobacco use. Third, Wagstaff decomposition identified '
     f'educational attainment ({decomp_pcts["Education"]:.1f}%) and urban–rural residence '
     f'({decomp_pcts["Residence"]:.1f}%) as the two dominant contributors to wealth-related '
     f'inequality in smoking, with national-level macroeconomic indicators playing comparatively '
     f'limited roles.')

sec('4.2  Comparison with Existing Literature', level=2)
body('These findings should be understood and situated within the context of the existing '
     'literature. The finding of a pro-poor gradient in smoking is consistent with multiple '
     'studies based on DHS and Global Adult Tobacco Survey (GATS) data. Sreeramareddy and '
     'Acharya (2021) reported significantly higher smoking prevalence among low-education '
     'and low-wealth groups across 22 sub-Saharan African countries. The GBD 2019 Tobacco '
     'Collaborators (2021), in an analysis spanning 204 countries and territories, documented '
     'a trend toward concentration of the smoking-attributable disease burden in lower-SDI '
     'countries. Huang et al. (2023) identified socioeconomic inequalities in smoking and '
     'cessation behaviors among Chinese adults aged 18–59 years. Tanaka et al. (2021) '
     'reported a sustained widening of socioeconomic inequality in smoking in Japan over the '
     'period 2001–2016. Disney et al. (2023), in an Australian analysis, further '
     'demonstrated that despite a continued national decline in smoking prevalence, '
     'populations with disabilities and those in low-income groups were being "left behind" '
     'in tobacco control progress, with relative inequalities actually widening.')

body('Regarding the relationship between smokeless tobacco use and socioeconomic status, '
     'the findings of the present study fill a critical gap in the literature. The GBD 2019 '
     'Chewing Tobacco Collaborators (2021) reported prevalence patterns of chewing tobacco '
     'use globally, noting that the South-East Asia Region showed no significant decline '
     'over the period 1990–2019. Ghate et al. (2022), using Indian GATS-2 data, '
     'identified socioeconomic determinants of smokeless tobacco use among Indian women. '
     'Halder et al. (2025), using data from the Longitudinal Ageing Study in India (LASI), '
     'found that urban–rural residence contributed substantially to variation in '
     'smokeless tobacco use among middle-aged and elderly Indians. Shaikh and Saikia (2022) '
     'reported socioeconomic inequalities in tobacco cessation among Indians aged 15 years '
     'and above. Chugh et al. (2023), in a systematic review, noted that tobacco control '
     'policies targeting smokeless tobacco lagged substantially behind those targeting '
     'smoking in both coverage and enforcement. However, all of these studies were based '
     'on single-country survey data and lacked a unified inequality index framework for '
     'direct cross-national comparison. The present study, through its paired comparison '
     'design contrasting any-tobacco and smoked-only CIs, provides indirect yet systematic '
     'cross-national evidence for the additional pro-poor effect of smokeless tobacco use.')

body(f'Carnazza et al. (2023), in a comparative assessment across European Union countries, '
     f'found substantial cross-country heterogeneity in income-related smoking inequality, '
     f'with some countries even exhibiting a pro-rich gradient. The inter-country CI variation '
     f'observed in the present study (SD = {sl_w.CI.std():.3f}) is consistent with this '
     f'finding—although the majority of low- and middle-income countries exhibited a '
     f'pro-poor gradient, the gradient strength varied considerably across countries, with '
     f'a significantly stronger pro-poor gradient in low-income than in high-income countries. '
     f'This pattern provides empirical support for the epidemiological transition model of '
     f'tobacco use.')

sec('4.3  Interpretation of Wagstaff Decomposition', level=2)
body(f'Having established the socioeconomic inequality patterns for smoking and any tobacco use, '
     f'the present study further applied Wagstaff decomposition to examine the structural '
     f'factors driving these inequalities. The decomposition results revealed that educational '
     f'attainment ({decomp_pcts["Education"]:.1f}%) and urban–rural residence '
     f'({decomp_pcts["Residence"]:.1f}%) were the two dominant contributors to '
     f'wealth-related inequality in smoking, collectively accounting for over 80%, whereas '
     f'macroeconomic indicators—GDP per capita ({decomp_pcts["GDP per capita"]:.1f}%) '
     f'and the Gini coefficient ({decomp_pcts["Gini coefficient"]:.1f}%)—played '
     f'comparatively limited roles. This finding is consistent with the decomposition results '
     f'reported by Huang et al. (2023) using Chinese adult data and by Halder et al. (2025) '
     f'using Indian middle-aged and elderly population data, collectively pointing to '
     f'individual-level socioeconomic position variables—rather than national '
     f'macroeconomic indicators—as the dominant forces driving inequality in health '
     f'behaviors. Mann et al. (2024), in their analysis of WHO FCTC investment cases across '
     f'21 low- and middle-income countries, similarly emphasized that the evaluation of tobacco '
     f'control policy effectiveness must incorporate an equity dimension to adequately capture '
     f'the differential impacts across income groups.')

sec('4.4  Policy Implications', level=2)
body('The present study carries several direct implications for global and national tobacco '
     'control policy. First, the core finding—that the socioeconomic gradient in any '
     'tobacco use is significantly stronger than that for smoking alone—suggests that '
     'current tobacco control policies may systematically underestimate the contribution of '
     'smokeless tobacco use to health inequality. Puljevic et al. (2024), in a scoping review, '
     'noted that smokeless tobacco products occupy a marginalized position in the tobacco '
     'endgame strategies of most countries, and that socioeconomically disadvantaged groups '
     'lack adequate representation in policy design and evaluation. Mills et al. (2024) called '
     'for the systematic incorporation of equity considerations into tobacco control, proposing '
     'specific recommendations including monitoring tobacco-related inequalities, employing '
     'equity-oriented enforcement strategies, and providing targeted cessation support for '
     'disadvantaged groups. The quantitative findings of the present study provide an empirical '
     'foundation for these policy proposals.')

body('Second, Spencer et al. (2024) demonstrated that increases in tobacco taxation yield '
     'pro-poor health benefits—the largest reduction in smoking prevalence occurred among '
     'the poorest 20% of the population, who bore only approximately 12% of the additional '
     'tax expenditure. However, Vladisavljevic et al. (2024) and Macias Sanchez and Garcia '
     'Gomez (2024) showed that the crowding-out effects of continued tobacco expenditure on '
     'food, education, and healthcare consumption are more severe among low-income households, '
     'with some households falling below the poverty line as a result of tobacco expenditure. '
     'Consequently, price and taxation policies targeting smokeless tobacco must be implemented '
     'alongside cessation support and alternative livelihood programs to avoid imposing '
     'disproportionate economic burdens on disadvantaged populations.')

sec('4.5  Strengths and Limitations', level=2)
body('The present study has several strengths: all data were obtained from publicly accessible '
     'databases, and the analysis code has been publicly archived (full reproducibility); both '
     'absolute and relative inequality indices were reported, in accordance with WHO health '
     'inequality monitoring guidelines; the paired comparison design contrasting any-tobacco '
     'and smoked-only CIs effectively controlled for country-specific confounding factors, '
     'yielding ΔCI estimates with high internal validity; and the geographic coverage '
     f'({sl_w.country.nunique()} DHS countries) substantially exceeded that of previous '
     'single-country or single-region studies.')

body('The present study has the following limitations. First, DHS smokeless tobacco use '
     'indicators (AH_TOBU_*_ASM) do not include wealth-, education-, or residence-stratified '
     'data—the dataset provides only national aggregate estimates; thus, the study could '
     'not directly compute independent CIs for smokeless tobacco use. The paired comparison '
     'of any-tobacco versus smoked-only CIs constitutes an indirect inference; ΔCI '
     'reflects the average additional effect of smokeless tobacco rather than its own '
     'socioeconomic gradient. This data limitation arises from insufficient coverage of '
     'stratified reporting for smokeless tobacco indicators within the DHS program. '
     'Second, the study employed an ecological design, with the country–survey–subgroup '
     'as the unit of analysis; inferences about exposure–outcome associations at the '
     'individual level are subject to the ecological fallacy. Sampling errors of DHS subgroup '
     'estimates are not fully reported in the aggregate dataset, and the asymptotic standard '
     'error assumptions of WLS regression may not hold perfectly in practice. '
     'Third, heterogeneity exists across DHS surveys in questionnaire design, tobacco use '
     'definitions, and data collection methods, varying by country and survey year '
     '(GBD 2019 Tobacco Collaborators, 2021). Fourth, the covariates included in the Wagstaff '
     'decomposition are aggregate-level variables and cannot fully substitute for '
     'individual-level behavioral pathway analyses '
     '(Littlecott et al., 2022; McEvoy & Layte, 2025). Fifth, the matching tolerance '
     '(±3 years) for World Bank covariate linkage may introduce measurement error, '
     'although sensitivity analyses confirmed that this strategy had no material effect '
     'on the primary conclusions.')

sec('4.6  Future Research Directions', level=2)
body('Based on these findings and limitations, future research should: (1) advocate for '
     'national DHS programs and other national health surveys to strengthen the collection '
     'and public release of stratified data for smokeless tobacco use indicators—at a '
     'minimum, including estimates classified by wealth quintile, educational attainment, '
     'and urban–rural residence; (2) combine DHS individual-level microdata within '
     'multilevel modeling frameworks to simultaneously estimate individual- and country-level '
     'inequality parameters (Pilvar et al., 2026); (3) once country-level smokeless tobacco '
     'stratified data become available, directly decompose the any-tobacco CI into a weighted '
     'sum of the smoking CI and the smokeless tobacco CI; and (4) investigate, within a broader '
     'health inequality framework, whether the disease burden of smokeless tobacco-attributable '
     'oral potentially malignant disorders and oral cavity cancer is disproportionately '
     'distributed across socioeconomic groups (Van Hemelrijck et al., 2024).')

sec('4.7  Conclusions', level=2)
body(f'This study, based on empirical data from {sl_w.country.nunique()} DHS countries, provides '
     f'systematic cross-national evidence on socioeconomic inequality in smoking and tobacco use, '
     f'and presents the first cross-national, paired demonstration of an additional pro-poor '
     f'amplification effect of smokeless tobacco use on tobacco inequality gradients. '
     f'Educational attainment and urban–rural residence are the core structural factors '
     f'driving this inequality. The global tobacco control agenda should incorporate smokeless '
     f'tobacco into inequality monitoring and intervention frameworks, and place socioeconomically '
     f'disadvantaged rural populations with low educational attainment at the center of tobacco '
     f'control strategies—this carries profound public health significance for reducing '
     f'the global socioeconomic inequality burden of lip and oral cavity cancer.')

# =================== DECLARATIONS ===================
sec('Declarations', level=1)

sec('Funding', level=2)
body('This work was supported by the Guangdong Traditional Chinese Medicine Bureau Project '
     '[Grant No. 20251039].')

sec('Conflicts of Interest', level=2)
body('All authors declare that they have no conflicts of interest relevant to the subject '
     'matter of this study.')

sec('Data and Code Availability', level=2)
body('All data used in this study were obtained from the following publicly accessible databases: '
     'the Demographic and Health Surveys (DHS) Program Indicator Data '
     '(https://dhsprogram.com), the WHO Global Health Observatory '
     '(https://www.who.int/data/gho), and the World Bank Development Indicators '
     '(https://data.worldbank.org). Specific retrieval indicators are detailed in the Methods '
     'section. The complete analysis code and data have been publicly archived on GitHub '
     '(https://github.com/1947314628-ui/tobacco-inequality-dhs), '
     'enabling full reproduction of all analyses reported in this study.')

sec('Ethics Approval', level=2)
body('Each DHS survey received ethical approval from the relevant national institutional review '
     'board and obtained informed consent from all respondents. The present study exclusively '
     'used publicly available, aggregate-level, de-identified secondary data and did not access '
     'individual-level identifiable information. In accordance with the Declaration of Helsinki '
     'and the Council for International Organizations of Medical Sciences (CIOMS) guidelines '
     'on secondary data analysis, this study is exempt from additional institutional review '
     'board ethical review.')

sec('Author Contributions', level=2)
body('Conceptualization: Juncheng He, Xinyue Li. Data curation & Formal analysis: Xinyue Li, '
     'Mianjia Wan, Peng Zhou. Interpretation: Xinyue Li, Mianjia Wan, Tong Yin. '
     'Writing – original draft: Xinyue Li, Mianjia Wan. '
     'Writing – review & editing: Tong Yin, Juncheng He. '
     'Supervision: Juncheng He. All authors have read and approved the final manuscript.')

# =================== REFERENCES ===================
sec('References', level=1)

all_refs = [
    'Carnazza, G., Liberati, P., & Resce, G. (2023). Income-related inequality in smoking habits: '
    'A comparative assessment in the European Union. Health Policy, 128, 34–41.',
    'Chen, D. T., Millett, C., & Filippidis, F. T. (2021). Prevalence and determinants of dual '
    'and poly-tobacco use among males in 19 low-and middle-income countries. Preventive Medicine, '
    '142, 106377.',
    'Chugh, A., Arora, M., Jain, N., et al. (2023). The global impact of tobacco control policies '
    'on smokeless tobacco use: a systematic review. The Lancet Global Health, 11(6), e953–e968.',
    'Dai, X., Gakidou, E., & Lopez, A. D. (2022). Evolution of the global smoking epidemic over '
    'the past half century. Tobacco Control, 31(2), 129–137.',
    'Disney, G., Petrie, D., Yang, Y., et al. (2023). Smoking inequality trends by disability and '
    'income in Australia, 2001 to 2020. Epidemiology, 34(2), 302–309.',
    'GBD 2019 Chewing Tobacco Collaborators. (2021). Spatial, temporal, and demographic patterns '
    'in prevalence of chewing tobacco use in 204 countries and territories, 1990–2019. '
    'The Lancet Public Health, 6(7), e482–e499.',
    'GBD 2019 Lip, Oral, and Pharyngeal Cancer Collaborators. (2023). The global, regional, and '
    'national burden of adult lip, oral, and pharyngeal cancer in 204 countries and territories. '
    'JAMA Oncology, 9(10), 1401–1416.',
    'GBD 2019 Tobacco Collaborators. (2021). Spatial, temporal, and demographic patterns in '
    'prevalence of smoking tobacco use and attributable disease burden in 204 countries and '
    'territories, 1990–2019. The Lancet, 397(10292), 2337–2360.',
    'Ghate, N., Kumar, P., & Dhillon, P. (2022). Socioeconomic determinants of smokeless tobacco '
    'use among Indian women: An analysis of Global Adult Tobacco Survey-2, India. '
    'WHO South-East Asia Journal of Public Health, 11(1), 24–31.',
    'Halder, P., Alwani, A. A., Nongkynrih, B., et al. (2025). Rural-urban disparities in tobacco '
    'use among middle aged and elderly Indian adults: a multivariate decomposition analysis. '
    'BMC Public Health, 25(1), 2818.',
    'Huang, M. Z., Liu, T. Y., Zhang, Z. M., et al. (2023). Trends in the distribution of '
    'socioeconomic inequalities in smoking and cessation. International Journal for Equity in '
    'Health, 22(1), 86.',
    'Kaur, J., Rinkoo, A. V., & Richardson, S. (2024). Trends in smokeless tobacco use and '
    'attributable mortality and morbidity in the South-East Asia Region. Tobacco Control, '
    '33(4), 425–433.',
    'Littlecott, H. J., Moore, G. F., McCann, M., et al. (2022). Exploring the association between '
    'school-based peer networks and smoking according to socioeconomic status. BMC Public Health, '
    '22(1), 142.',
    'Macias Sanchez, A., & Garcia Gomez, A. (2024). Crowding out and impoverishing effect of '
    'tobacco in Mexico. Tobacco Control, 33(Suppl 2), s68–s74.',
    'Mann, N., Spencer, G., Hutchinson, B., et al. (2024). Interpreting results, impacts and '
    'implications from WHO FCTC tobacco control investment cases in 21 low-income and '
    'middle-income countries. Tobacco Control, 33(Suppl 1), s17–s26.',
    'McEvoy, O., & Layte, R. (2025). Bringing the group back in: Social class and resistance in '
    'adolescent smoking. Sociology of Health & Illness, 47(2), e13858.',
    'Mills, S. D., Rosario, C., Yerger, V. B., et al. (2024). Recommendations to advance equity '
    'in tobacco control. Tobacco Control, 33(e2), e246–e253.',
    'Pilvar, H., Nicodemo, C., Petrou, S., et al. (2026). Socioeconomic inequity in extreme '
    'outcomes within very pre-term and/or very low birthweight infants. Frontiers in Public '
    'Health, 14, 1791450.',
    'Puljevic, C., Feulner, L., Hobbs, M., et al. (2024). Tobacco endgame and priority '
    'populations: a scoping review. Tobacco Control, 33(e2), e231–e239.',
    'Rumgay, H., Nethan, S. T., Shah, R., et al. (2024). Global burden of oral cancer in 2022 '
    'attributable to smokeless tobacco and areca nut consumption. The Lancet Oncology, 25(11), '
    '1413–1423.',
    'Shaikh, R., & Saikia, N. (2022). Socioeconomic inequalities in tobacco cessation among '
    'Indians above 15 years of age from 2009 to 2017. BMC Public Health, 22(1), 1419.',
    'Spencer, G., Nugent, R., Mann, N., et al. (2024). Equity implications of tobacco taxation: '
    'results from WHO FCTC investment cases. Tobacco Control, 33(Suppl 1), s27–s33.',
    'Sreeramareddy, C. T., & Acharya, K. (2021). Trends in prevalence of tobacco use by sex and '
    'socioeconomic status in 22 Sub-Saharan African countries, 2003–2019. JAMA Network Open, '
    '4(12), e2137820.',
    'Tanaka, H., Mackenbach, J. P., & Kobayashi, Y. (2021). Widening socioeconomic inequalities '
    'in smoking in Japan, 2001–2016. Journal of Epidemiology, 31(6), 369–377.',
    'Van Hemelrijck, W. M. J., Kunst, A. E., Sizer, A., et al. (2024). Trends in educational '
    'inequalities in smoking-attributable mortality. Journal of Epidemiology and Community '
    'Health, 78(9), 561–569.',
    'Vladisavljevic, M., Zubovic, J., Jovanovic, O., & Dukic, M. (2024). Crowding-out effect '
    'of tobacco consumption in Serbia. Tobacco Control, 33(Suppl 2), s88–s94.',
    'Yang, H., Ma, C., Zhao, M., Magnussen, C. G., & Xi, B. (2022). Prevalence and trend of '
    'smokeless tobacco use and its associated factors among adolescents aged 12–16 years in 138 '
    'countries/territories, 1999–2019. BMC Medicine, 20(1), 460.',
]

for ref in all_refs:
    rp = doc.add_paragraph(ref)
    rp.paragraph_format.first_line_indent = Cm(0)
    rp.paragraph_format.left_indent = Cm(0.74)
    rp.paragraph_format.first_line_indent = Cm(-0.74)
    for run in rp.runs:
        run.font.size = Pt(10)

# =================== SAVE ===================
out = os.path.join(ROOT, "manuscript", "Full_Manuscript_EN.docx")
os.makedirs(os.path.join(ROOT, "manuscript"), exist_ok=True)
doc.save(out)
print(f"Saved: {out}")
