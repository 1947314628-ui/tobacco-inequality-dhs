#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate Supplementary Materials Word document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os, pandas as pd, numpy as np

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

res = pd.read_csv(os.path.join(ROOT, "data", "inequality_results.csv"))
master = pd.read_csv(os.path.join(ROOT, "data", "analysis_master.csv"))
gho = pd.read_csv(os.path.join(ROOT, "data", "gho_tidy.csv"))

doc = Document()
s = doc.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(11)
s.paragraph_format.line_spacing = 1.5

for i in range(1, 3):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'; hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.size = Pt(14) if i == 1 else Pt(12); hs.font.bold = True


def add_tbl(doc, headers, rows, caption):
    pc = doc.add_paragraph(); pc.paragraph_format.space_before = Pt(10)
    r = pc.add_run(caption); r.font.size = Pt(10); r.font.bold = True
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Light Shading Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for rn in p.runs: rn.font.size = Pt(9); rn.font.bold = True
    for i, rd in enumerate(rows):
        for j, v in enumerate(rd):
            c = t.rows[i + 1].cells[j]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for rn in p.runs: rn.font.size = Pt(9)
    return t


title = doc.add_heading('Supplementary Materials', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# -- data prep --
dhs_sm = master[(master.data_source == 'DHS') & (master.indicator_abbr == 'smoked')
                & (master.dimension.isin(['wealth', 'education', 'residence']))]
smoke_prev = dhs_sm.groupby('country').estimate.mean()

sl_w = res[(res.indicator_abbr == 'smoked') & (res.dimension == 'wealth')].dropna(subset=['CI'])
gho_sl = gho[(gho.indicator_abbr == 'smokeless') & (gho.subgroup == 'SEX_BTSX')].dropna(subset=['estimate'])
top15_sl = gho_sl.nlargest(15, 'estimate')

latest_idx = sl_w.groupby('iso3').year.idxmax()

a = res[(res.indicator_abbr == 'any_tobacco') & (res.dimension == 'wealth')][['iso3', 'country', 'CI', 'CI_se']]
a = a.rename(columns={'CI': 'CI_any', 'CI_se': 'CI_se_any'})
b = res[(res.indicator_abbr == 'smoked') & (res.dimension == 'wealth')][['iso3', 'CI', 'CI_se']]
b = b.rename(columns={'CI': 'CI_sm', 'CI_se': 'CI_se_sm'})
paired = a.merge(b, on='iso3').dropna(); paired['delta'] = paired.CI_any - paired.CI_sm

# ===== Table S1 =====
doc.add_heading('Table S1. Descriptive statistics of DHS countries included in the analysis', level=2)
t1_rows = []
for ctry in smoke_prev.sort_values(ascending=False).index:
    sub = dhs_sm[dhs_sm.country == ctry]
    gdp_val = f'{sub.gdp_pc_usd.mean():.0f}' if 'gdp_pc_usd' in sub.columns and sub.gdp_pc_usd.notna().any() else 'N/A'
    t1_rows.append([ctry, str(int(sub.year.nunique())), f'{sub.estimate.mean():.1f}',
                    f'{sub.estimate.min():.1f}-{sub.estimate.max():.1f}',
                    str(int(sub.year.max())), gdp_val])
add_tbl(doc, ['Country', 'N surveys', 'Mean prevalence (%)', 'Subgroup range (%)',
              'Latest year', 'GDP per capita (USD)'], t1_rows, '')

# ===== Table S2 =====
doc.add_heading('Table S2. Wealth-related inequality indices for tobacco smoking (all country-survey observations)', level=2)
t2_rows = []
for _, r in sl_w.sort_values('CI').iterrows():
    t2_rows.append([r['country'], str(int(r['year'])), f'{r["CI"]:.4f}',
                    f'{r["CI_se"]:.4f}', f'{r["SII"]:.2f}', f'{r["RII"]:.2f}'])
add_tbl(doc, ['Country', 'Year', 'CI', 'SE', 'SII (pp)', 'RII'], t2_rows, '')

# ===== Table S3 =====
doc.add_heading('Table S3. Top 15 countries by adult current smokeless tobacco use prevalence (WHO GHO)', level=2)
t3_rows = [[r['country'], f'{r["estimate"]:.1f}', str(int(r['year']))] for _, r in top15_sl.iterrows()]
add_tbl(doc, ['Country', 'Prevalence (%)', 'Year'], t3_rows, '')

# ===== Table S4 =====
doc.add_heading('Table S4. Sensitivity analysis of wealth-related CI for tobacco smoking', level=2)
s_rows = [
    ['All surveys', str(len(sl_w)), f'{sl_w.CI.mean():.4f}', f'{sl_w.SII.mean():.2f}'],
    ['Latest survey per country', str(len(latest_idx)), f'{sl_w.loc[latest_idx].CI.mean():.4f}', f'{sl_w.loc[latest_idx].SII.mean():.2f}'],
    ['Surveys 2015 and later', str(len(sl_w[sl_w.year >= 2015])), f'{sl_w[sl_w.year >= 2015].CI.mean():.4f}', f'{sl_w[sl_w.year >= 2015].SII.mean():.2f}'],
    ['Surveys before 2015', str(len(sl_w[sl_w.year < 2015])), f'{sl_w[sl_w.year < 2015].CI.mean():.4f}', f'{sl_w[sl_w.year < 2015].SII.mean():.2f}'],
]
add_tbl(doc, ['Scenario', 'N', 'Mean CI', 'Mean SII (pp)'], s_rows, '')

# ===== Table S5 =====
doc.add_heading('Table S5. Paired comparison of wealth-related CI: all-tobacco use vs. smoked tobacco only', level=2)
t5_rows = []
for _, r in paired.sort_values('delta').iterrows():
    t5_rows.append([r['country'], f'{r["CI_any"]:.4f}', f'{r["CI_sm"]:.4f}', f'{r["delta"]:.4f}'])
add_tbl(doc, ['Country', 'CI_all_tobacco', 'CI_smoked', 'Delta CI'], t5_rows, '')

out = os.path.join(ROOT, "supplementary", "Supplementary_Materials.docx")
os.makedirs(os.path.join(ROOT, "supplementary"), exist_ok=True)
doc.save(out)
print(f"Saved: {out}")
