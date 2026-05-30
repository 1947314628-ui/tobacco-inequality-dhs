#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate concise Methods section — journal format, ~3-4 pages."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
doc = Document()

# --- Styles ---
s = doc.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 2.0
s.paragraph_format.space_after = Pt(0)
s.paragraph_format.first_line_indent = Cm(0.74)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'; hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.first_line_indent = Cm(0)
    hs.paragraph_format.space_before = Pt(10); hs.paragraph_format.space_after = Pt(4)
    hs.font.size = [13, 12, 12][i - 1]; hs.font.bold = True


def add_tbl(doc, headers, rows, caption):
    pc = doc.add_paragraph(); pc.paragraph_format.first_line_indent = Cm(0)
    pc.paragraph_format.space_before = Pt(8)
    r = pc.add_run(caption); r.font.size = Pt(10); r.font.bold = True
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Light Shading Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for rn in p.runs: rn.font.size = Pt(9); rn.font.bold = True
    for i, rd in enumerate(rows):
        for j, v in enumerate(rd):
            c = t.rows[i + 1].cells[j]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for rn in p.runs: rn.font.size = Pt(9)


def add_flowchart(doc):
    pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.first_line_indent = Cm(0)
    pc.paragraph_format.space_before = Pt(8)
    r = pc.add_run('图1  研究筛选流程图'); r.font.size = Pt(10); r.font.bold = True

    items = [
        ["DHS指标API检索\n(AH_TOBC_*, AH_TOBU_*, AH_TOBA_*; breakdown=all)", "15,677条"],
        ["↓ 排除: 非目标维度(Total/Age/Region等)", "−12,293"],
        ["具有财富/教育/城乡分层数据", "3,384条"],
        ["↓ 排除: 亚组排序缺失或分类不完整", "−1,208"],
        ["最终DHS分析样本", "2,176条"],
        [" ├ 吸烟×财富(5 quintiles)", "178观测, 66国"],
        [" ├ 吸烟×教育(4 levels)", "169观测, 61国"],
        [" ├ 吸烟×城乡(2 categories)", "178观测, 66国"],
        [" └ 全烟草×财富(5 quintiles)", "21观测, 21国"],
        ["", ""],
        ["WHO GHO: Adult_curr_smokeless (≥2010)", "277条, 102国"],
        ["WHO GHO: M_Est_smk_curr_std (≥2010)", "4,008条, 167国"],
        ["世界银行: 5项宏观指标 (2010–2024)", "4,827条, 261国"],
        [" 协变量完整匹配DHS观测", "61国 (92.4%)"],
    ]
    tbl = doc.add_table(rows=len(items), cols=2)
    tbl.style = 'Light Grid Accent 1'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (left, right) in enumerate(items):
        c0 = tbl.rows[i].cells[0]; c0.text = left
        c1 = tbl.rows[i].cells[1]; c1.text = right
        for c in [c0, c1]:
            for p in c.paragraphs:
                for rn in p.runs: rn.font.size = Pt(8)
    for row in tbl.rows:
        row.cells[0].width = Cm(8); row.cells[1].width = Cm(8)


# ================================================================
t = doc.add_heading('方法', level=1)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ================================================================
doc.add_heading('2.1 研究设计与数据来源', level=2)

doc.add_paragraph(
    '本研究为横断面生态学分析，分析单元为国家-调查观测。数据来源于三个公开数据库：'
    '(1) 人口与健康调查项目（The DHS Program, https://dhsprogram.com）指标数据，'
    'DHS为两阶段分层整群抽样设计的全国代表性住户调查，本研究使用其发布的聚合指标估计值，'
    '检索指标包括AH_TOBC_W_ANY、AH_TOBC_M_ANY（当前吸烟）、AH_TOBU_W_ASM、'
    'AH_TOBU_M_ASM（当前无烟烟草使用）及AH_TOBA_W_ANY、AH_TOBA_M_ANY（当前任何烟草使用），'
    '检索日期为2026年5月28日；'
    '(2) 世界卫生组织全球卫生观察站（WHO GHO, https://www.who.int/data/gho），'
    '检索Adult_curr_smokeless（成人当前无烟烟草使用）及M_Est_smk_curr_std '
    '（年龄标准化当前吸烟现患率）两项指标，限定空间维度为国家级、时间维度≥2010年；'
    '(3) 世界银行发展指标（World Bank, https://data.worldbank.org），'
    '检索人均GDP（NY.GDP.PCAP.CD）、基尼系数（SI.POV.GINI）、城镇人口占比'
    '（SP.URB.TOTL.IN.ZS）、卫生支出占GDP百分比（SH.XPD.CHEX.GD.ZS）及15岁以上'
    '人口平均受教育年限（BAR.SCHL.15UP），限定年份2010–2024年。'
    '世界银行指标通过ISO 3166-1 alpha-3国家代码与DHS观测关联，匹配容忍度为调查年份±3年。'
    '全部数据为聚合层面、已去标识化的公开二手数据。'
)

# ================================================================
doc.add_heading('2.2 研究人群与筛选流程', level=2)

doc.add_paragraph(
    'DHS检索参数设定为breakdown = all（获取所有亚组分类）及IsPreferred = 1'
    '（DHS推荐的优先估计值），初始返回15,677条记录。纳入标准为：'
    '(1) CharacteristicCategory属于"Wealth quintile""Education"或"Residence"'
    '（即具备财富五分位、教育水平或城乡居住地的社会经济分层）；'
    '(2) 亚组标签具备明确的有序分类（如财富最低至最高、教育从无至高等）；'
    '(3) 同一国家-调查内，该维度全部亚组均无缺失（如财富维度须具备全部五个五分位）。'
    '排除标准为：CharacteristicCategory为"Total"（仅全国总计）、'
    '"Age"分组、"Region"分组等非社会经济维度的记录。'
)

doc.add_paragraph(
    '经筛选（图1），最终纳入吸烟财富五分位分层数据178组国家-调查观测（66个国家）、'
    '教育分层169组（61国）、城乡分层178组（66国）。全烟草使用（含无烟烟草制品）'
    '财富分层数据21组（21国）用于配对比较。男女分层数据通过各亚组的加权人口分母'
    '（DenominatorWeighted）加权合并为男女合计估计值。无烟烟草指标（AH_TOBU_*_ASM）'
    '经相同筛选后仅含有全国总计数据，不具备财富、教育或城乡分层信息，该指标仅用于'
    '描述国家层面现患率，不纳入不平等指数计算。'
)

add_flowchart(doc)

# ================================================================
doc.add_heading('2.3 变量定义', level=2)

doc.add_paragraph(
    '暴露变量（不平等维度）：(1) 财富五分位（Wealth quintile），DHS财富指数按'
    '家庭五分位分为Lowest (order=1)至Highest (order=5)；(2) 教育水平（Education），'
    '分为No education (order=1)、Primary (2)、Secondary (3)、Higher (4)；'
    '(3) 城乡居住地（Residence），分为Rural (order=1)、Urban (order=2)。'
    '各维度中order值较小表示社会经济弱势程度较高。'
)

doc.add_paragraph(
    '结局变量为各亚组的加权现患率（%），即DHS返回的Value字段值（该亚组当前吸烟者'
    '或任何烟草使用者占15–49岁人口的百分比），为连续型比例变量（0–100%）。'
    'DHS发布的估计值已经过各调查的住户抽样权重校正，本研究在聚合分析中使用这些'
    '加权估计值及其加权人口分母（DenominatorWeighted）作为亚组人口权重。'
)

doc.add_paragraph(
    '国家层面协变量：对数人均GDP（log-GDP，现价美元）、基尼系数（0–100）、'
    '城镇人口占比（%）、卫生支出占GDP百分比（%）。各变量为连续型，'
    '通过国家代码及最近调查年份匹配至DHS观测。'
)

# ================================================================
doc.add_heading('2.4 缺失数据', level=2)

doc.add_paragraph(
    'DHS API层面仅返回具有有效估计值的记录，不存在亚组内估计值缺失。'
    '在66个DHS国家中，61国（92.4%）具有全部四项世界银行协变量的完整数据；'
    '平均受教育年限（BAR.SCHL.15UP）在多数国家-年份组合中缺失（144/4,827条记录非缺失），'
    '未纳入主要分析。Wagstaff分解限定于协变量完整的61个国家（完整案例分析），'
    '未实施数据插补。主要不平等指数（CI、SII、RII）的计算不依赖于世界银行协变量，'
    '因此不受协变量缺失影响。敏感性分析中比较了年份匹配容差±3年与±5年的结果。'
)

# ================================================================
doc.add_heading('2.5 统计分析', level=2)

doc.add_heading('2.5.1 复杂调查设计', level=3)
doc.add_paragraph(
    'DHS各次调查均采用两阶段分层整群抽样设计，DHS发布的聚合指标估计值已在计算阶段'
    '施加了抽样权重校正。本研究的分析单元为聚合层面（国家-调查-亚组），'
    '不平等指数计算以各亚组的加权人口份额（DenominatorWeighted）为权重。'
    'CI及SII的标准误通过加权最小二乘回归的渐近理论推导得出。'
    '本研究未涉及个体层面微数据的直接建模，因此未使用专门的调查数据分析模块。'
)

doc.add_heading('2.5.2 不平等指数', level=3)
doc.add_paragraph(
    '对每个国家-调查-维度组合，亚组按order升序排列（最弱势→最有利），'
    '以人口份额为权重计算ridit秩（R_i = Σ_{j=1}^{i-1} w_j + w_i/2）。'
    '斜率不平等指数（SII）通过加权最小二乘回归estimate_i = β_0 + β_1 × R_i + ε_i '
    '估计，SII = β_1（单位为百分点），SII < 0表示现患率集中于弱势群体。'
    '相对不平等指数RII = ŷ(R=1) / ŷ(R=0)。'
    '集中指数（CI）采用便捷回归法：CI = (2/μ) × Σ w_i × estimate_i × (R_i − 0.5)，'
    '取值范围(−1, +1)，负值表示集中于弱势群体。CI标准误通过'
    'yi* = 2·var(R)·estimate_i/μ对R_i的WLS回归系数标准误得出。'
    'Wagstaff归一化CI = CI/(1 − p)用于校正现患率比例的有界性'
    '（p = μ/100，若μ > 1）。所有指标均报告点估计值及95%置信区间。'
)

doc.add_heading('2.5.3 分解分析', level=3)
doc.add_paragraph(
    '在61个协变量完整的国家中，对吸烟财富相关CI实施Wagstaff分解。'
    '解释变量包括教育CI、城乡CI、标准化对数人均GDP、标准化基尼系数、标准化城镇人口占比'
    '及标准化卫生支出，拟合多元线性回归模型：CI_wealth = α + Σ β_k × X_k + ε。'
    '各因子相对贡献以|β_k × SD(X_k)| / Σ|β_j × SD(X_j)| × 100%表示。'
    '模型R²为可解释变异比例，1−R²为残差成分。'
)

doc.add_heading('2.5.4 分层与敏感性分析', level=3)
doc.add_paragraph(
    '预设分层分析：按人均GDP四分位数（低收入/中低收入/中高收入/高收入）及调查时期'
    '（2010–2014年 vs. 2015年及以后）分层，各层报告合并均值CI及95% CI。'
    '全烟草与纯吸烟CI的配对比较采用配对t检验，报告均值ΔCI及95% CI。'
    '敏感性分析：(1) 以Wagstaff归一化CI替代标准CI；(2) 仅保留每国最近一次调查；'
    '(3) 按调查时期分层；(4) 协变量匹配容差扩展至±5年。'
)

doc.add_heading('2.5.5 分析软件', level=3)
doc.add_paragraph(
    '数据分析在Python 3.12环境中完成，使用pandas (v2.0+)、numpy (v1.24+)、'
    'scipy (v1.10+)及scikit-learn (v1.3+)进行统计计算，matplotlib (v3.7+)出图。'
    '全部代码已公开存档（见数据与代码可及性声明）。'
)

# ================================================================
doc.add_heading('2.6 伦理声明', level=2)

doc.add_paragraph(
    'DHS项目各次调查均已获得各国相应伦理审查委员会批准及受访者知情同意。'
    '本研究仅使用公开可获取的、聚合层面的、已去标识化的二手数据，'
    '未访问个体层面标识信息。根据《赫尔辛基宣言》及国际医学科学组织理事会'
    '（CIOMS）指南，此类公开聚合数据的二次分析免于机构审查委员会额外伦理审查。'
)

# ================================================================
out = os.path.join(ROOT, "manuscript", "Methods_CHS.docx")
os.makedirs(os.path.join(ROOT, "manuscript"), exist_ok=True)
doc.save(out)
print(f"Saved: {out}")
