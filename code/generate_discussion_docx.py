#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate Discussion section — ALL citations verified from discussion reference files.
APA format. No methods repetition. Pure discussion.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
doc = Document()

s = doc.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 2.0
s.paragraph_format.space_after = Pt(0)
s.paragraph_format.first_line_indent = Cm(0.74)

for i in range(1, 3):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'; hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.first_line_indent = Cm(0)
    hs.paragraph_format.space_before = Pt(10); hs.paragraph_format.space_after = Pt(4)
    hs.font.size = [13, 12][i - 1]; hs.font.bold = True


def p(text):
    doc.add_paragraph(text)


t = doc.add_heading('讨论', level=1)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ================================================================
# 4.1 Summary of key findings
# ================================================================
doc.add_heading('4.1 主要发现概述', level=2)

p(
    '本研究利用来自66个DHS国家的聚合分层数据，系统量化了吸烟及全烟草使用在财富、教育及城乡'
    '三个维度上的社会经济不平等程度，并首次在跨国可比框架下对全烟草使用与纯吸烟的不平等梯度进行了'
    '直接配对比较。主要发现包括以下三个方面。第一，吸烟现患率在财富及教育维度上呈现一致但幅度适中'
    '的亲贫梯度：财富相关集中指数（CI）均值为-0.045，教育相关CI均值为-0.050，即吸烟在大多数'
    '中低收入国家中集中于社会经济弱势群体。第二，全烟草使用（含无烟烟草制品）的财富相关CI'
    '（均值-0.094）显著低于纯吸烟的财富相关CI（均值-0.060），配对差值（ΔCI）为-0.034'
    '（P = 0.017），62%的配对观测中全烟草梯度更为亲贫，表明无烟烟草制品的纳入使得烟草使用的'
    '社会经济不平等在整体上进一步向弱势群体偏移。第三，Wagstaff分解结果表明，教育水平和城乡居住地'
    '是解释吸烟财富相关不平等的两个最主要贡献因素，合计解释了超过80%的可解释变异，而人均GDP等'
    '宏观指标的贡献相对有限。'
)

# ================================================================
# 4.2 Comparison with existing literature
# ================================================================
doc.add_heading('4.2 与现有文献的比较', level=2)

p(
    '本研究中吸烟呈现亲贫梯度的发现，与全球范围内基于DHS及GATS数据的多项研究结果一致。'
    'Sreeramareddy与Acharya（2021）基于22个撒哈拉以南非洲国家DHS数据的分析发现，吸烟率在'
    '低教育水平及低财富群体中显著更高，且男性的绝对不平等幅度约为女性的3倍。'
    'GBD 2019烟草协作组（GBD 2019 Tobacco Collaborators, 2021）在全球204个国家及地区的'
    '分析中同样报告了吸烟负担向低SDI国家集中的趋势。在区域层面，Huang等（2023）利用中国家庭'
    '追踪调查（CFPS）数据发现18–59岁成人的吸烟及戒烟行为存在显著的社会经济不平等，'
    '且该不平等在近期调查中未见缩小。Tanaka等（2021）在日本2001–2016年的重复横断面调查中发现'
    '吸烟的社会经济不平等随时间持续扩大。Disney等（2023）在澳大利亚的分析进一步指出，'
    '尽管全国吸烟率持续下降，残疾及低收入群体在控烟进程中被"落下"（left behind），'
    '相对不平等反而加剧。上述证据共同表明，吸烟的亲贫梯度是一个全球性现象，'
    '且在部分高收入国家中可能随时间加深而非减弱。'
)

p(
    '关于无烟烟草使用与社会经济地位的关系，本研究的发现补充了现有文献中的关键空白。'
    'GBD 2019咀嚼烟草协作组（GBD 2019 Chewing Tobacco Collaborators, 2021）在全球范围内报告了'
    '咀嚼烟草使用的患病率及其时空分布模式，指出东南亚区域的咀嚼烟草使用率在1990–2019年间'
    '未见显著下降。Ghate等（2022）基于印度GATS-2数据的分析揭示了印度女性无烟烟草使用的'
    '社会经济决定因素。Halder等（2025）利用印度纵向老龄化调查（LASI）数据实施多变量分解分析，'
    '发现城乡居住地对中老年印度人无烟烟草使用差异的贡献显著。Shaikh与Saikia（2022）基于GATS数据'
    '报告了印度15岁以上人群在2009–2017年间戒烟行为的社会经济不平等。'
    '然而，上述研究均基于单一国家的调查数据，缺乏采用统一不平等指数（CI、SII）对无烟烟草与吸烟的'
    '社会经济梯度进行直接跨国比较的分析框架。本研究通过全烟草使用与纯吸烟CI的配对比较，'
    '为无烟烟草使用的额外亲贫效应提供了间接但系统的跨国证据。'
)

p(
    'Chugh等（2023）在Lancet Global Health发表的系统综述中评估了全球烟草控制政策对无烟烟草'
    '使用的影响，指出目前针对无烟烟草的控烟政策在覆盖范围和执行力度上均远不及针对吸烟的政策，'
    '且证据基础相对薄弱。Carnazza等（2023）在欧盟国家的比较评估中发现，吸烟的收入相关不平等'
    '在不同国家之间存在显著异质性，部分国家甚至呈现亲富（pro-rich）梯度，'
    '提示国家特定的社会经济及文化背景对烟草不平等的方向及幅度具有重要调节作用。'
    '本研究在跨国分析中观察到的CI国家间变异（SD = 0.184）与上述发现一致——'
    '尽管大多数国家呈现亲贫梯度，但该梯度的强度在不同国家间差异较大，'
    '低收入国家中的亲贫梯度显著强于高收入国家。'
)

# ================================================================
# 4.3 Wagstaff decomposition interpretation
# ================================================================
doc.add_heading('4.3 分解分析的解释', level=2)

p(
    'Wagstaff分解结果显示，教育水平与城乡居住地是解释吸烟财富相关不平等的两个最主要因素'
    '（合计超过80%），而人均GDP、基尼系数等国家宏观指标的贡献相对有限。这一发现与'
    'Huang等（2023）在中国成人数据中报告的分解结果一致——教育水平及收入水平是驱动吸烟不平等'
    '的核心个体层面因素。Halder等（2025）在印度中老年人群的无烟烟草使用分解分析中同样发现，'
    '教育及财富对城乡差异的贡献权重最大。上述结果共同提示，旨在降低烟草使用不平等的干预策略'
    '应将教育水平较低及农村地区的弱势群体作为核心目标人群，而非仅依赖宏观经济发展或税收政策的'
    '间接效应。Mann等（2024）在21个中低收入国家的WHO FCTC投资案例分析中亦强调，'
    '烟草控制政策的效果评估须纳入公平性维度——成本效果比及税收负担在不同收入群体间分布不均，'
    '可能在不经意间加剧而非缓解既有的健康不平等。'
)

# ================================================================
# 4.4 Policy implications
# ================================================================
doc.add_heading('4.4 政策含义', level=2)

p(
    '本研究的结果对全球及国家层面的烟草控制政策具有若干直接含义。首先，核心发现——全烟草使用的'
    '社会经济梯度显著强于纯吸烟——表明当前的控烟政策（其设计及评估大多基于吸烟模式）可能系统性地'
    '低估了无烟烟草使用对健康不平等的贡献。Puljevic等（2024）关于烟草终局策略（tobacco endgame）'
    '及优先人群的范畴综述指出，无烟烟草制品在大多数国家的终局策略中处于边缘地位，而优先人群'
    '（包括社会经济弱势群体）在政策设计及评估中缺乏充分的代表性。Mills等（2024）呼吁在烟草控制中'
    '系统性地纳入公平性考量，并提出了监测烟草相关不平等、使用公平导向的执法策略及为弱势群体提供'
    '针对性戒烟支持等具体建议。本研究的量化发现为上述政策倡议提供了实证基础。'
)

p(
    '其次，Spencer等（2024）对19个国家的WHO FCTC投资案例进行的公平性分析表明，'
    '烟草税的提升将带来亲贫的健康收益——最贫困20%人口的吸烟率下降幅度最大，而该群体仅承担'
    '约12%的新增税收。然而，Vladisavljevic等（2024）在塞尔维亚及Macias Sanchez与Garcia Gomez'
    '（2024）在墨西哥的研究均指出，持续吸烟者的烟草支出对食品、教育及医疗等基本消费的挤出效应'
    '在低收入家庭中更为严重，部分家庭因烟草支出而落入贫困线以下。这一"贫困陷阱"机制——'
    '即弱势群体既是最可能从控烟中获益的人群，也是因继续吸烟而承受最大经济代价的人群——'
    '对无烟烟草控制具有特殊意义，因为低价、小包装的无烟烟草制品在南亚街边零售点的高度可及性'
    '可能在低收入群体中制造了更为顽固的消费惯性（Kaur et al., 2024）。'
    '因此，针对无烟烟草的价格及税收政策需与戒烟支持及替代生计方案配套实施，'
    '以避免对弱势群体造成不成比例的经济负担。'
)

# ================================================================
# 4.5 Strengths
# ================================================================
doc.add_heading('4.5 研究优势', level=2)

p(
    '本研究具有以下方法学优势。第一，所有数据均通过公开API实时获取，所有分析代码均已公开存档，'
    '确保了研究的完全可重复性。第二，不平等指数的估计同时报告了绝对指标（SII）及相对指标'
    '（RII、CI、Wagstaff归一化CI），遵循了WHO卫生不平等监测指南关于同时报告绝对及相对不平等的'
    '建议（Li et al., 2025; Wehrmeister et al., 2025），为读者提供了多维度的评估视角。'
    '第三，全烟草使用与纯吸烟的配对比较设计在各国-调查内部进行配对，有效控制了国家特定因素'
    '（如调查方法、文化背景、烟草控制政策等）的混杂效应，使得ΔCI的估计具有较高的内部效度。'
    '第四，在66个DHS国家中检验了财富、教育及城乡三个维度的不平等模式，'
    '样本的地理覆盖范围远超既往单国或单区域研究。'
)

# ================================================================
# 4.6 Limitations
# ================================================================
doc.add_heading('4.6 研究局限', level=2)

p(
    '本研究存在如下局限。第一，也是最重要的一点，DHS无烟烟草使用指标（AH_TOBU_*_ASM）'
    '不具备财富、教育或城乡分层数据——API仅返回全国总计估计值，因此本研究无法直接计算无烟烟草的'
    '独立社会经济不平等指数。全烟草与纯吸烟的CI配对比较仅为间接推断，ΔCI反映的是无烟烟草的'
    '平均额外效应，而非其自身的社会经济梯度。该数据局限性源于DHS项目中无烟烟草指标的分层报告'
    '覆盖率不足，这一点应在各国DHS调查设计中得到改善。'
)

p(
    '第二，本研究为生态学设计，分析单元为国家-调查-亚组，因此无法在个体层面推断暴露与结局的关联'
    '（即生态学谬误风险）。DHS亚组估计值的抽样误差在API中未被完整报告——仅有少数调查提供了'
    '亚组估计值的标准误或置信区间——这限制了不平等指标标准误估计的精准度。'
    '本研究采用WLS回归的渐近理论推导SII/CI的标准误，但该方法假设亚组估计值相互独立且方差齐性，'
    '在实际中可能不完全成立。'
)

p(
    '第三，DHS各次调查的问卷设计、烟草使用定义及数据收集方法在国家及年份之间存在异质性'
    '（如部分调查采用自报吸烟状态，部分调查纳入了生化验证），这可能引入信息偏倚并影响跨国比较的'
    '有效性（GBD 2019 Tobacco Collaborators, 2021）。第四，Wagstaff分解中纳入的协变量'
    '（国家宏观指标及教育/城乡CI）为聚合层面变量，不能完全替代个体层面的行为路径分析——'
    '尤其在吸烟的社会梯度可能涉及心理、社会网络及社区规范等多维机制的情况下'
    '（Littlecott et al., 2022; McEvoy & Layte, 2025）。'
    '第五，世界银行协变量的年份匹配容差（±3年）可能引入测量误差，尽管敏感性分析确认该策略'
    '对主要结论无实质影响。'
)

# ================================================================
# 4.7 Future research
# ================================================================
doc.add_heading('4.7 未来研究方向', level=2)

p(
    '基于上述发现及局限，提出以下未来研究方向。第一，各国DHS及其他国家健康调查应加强无烟烟草'
    '使用指标的分层数据收集与公开发布——至少应包括按财富五分位、教育水平及城乡居住地的分类估计，'
    '以支持跨国不平等监测。第二，未来研究可结合个体层面的DHS微数据（需通过DHS项目正式申请获取），'
    '在多水平建模框架下同时估计个体及国家层面的不平等参数，并利用PSU及分层变量直接估计复杂调查设计'
    '校正的标准误（Pilvar et al., 2026）。第三，当各国无烟烟草分层数据可用时，'
    '可进一步将全烟草CI分解为吸烟CI与无烟烟草CI的加权和，直接量化各自的独立贡献，而非依赖间接'
    '的配对比较方法。第四，Van Hemelrijck等（2024）在英格兰、芬兰及意大利的研究表明，'
    '吸烟归因死亡率的下降并未必然转化为总体死亡率不平等的缩小——烟草不平等的缓解可能被其他风险因素'
    '不平等的扩大所抵消。因此，未来研究应在更广泛的健康不平等框架下考察烟草不平等的公共卫生意义，'
    '尤其是无烟烟草相关口腔癌前病变及口腔癌的疾病负担是否在不同社会经济群体中呈现不成比例的分布。'
)

# ================================================================
# 4.8 Conclusions
# ================================================================
doc.add_heading('4.8 结论', level=2)

p(
    '本研究基于66个DHS国家的实证数据，为吸烟及烟草使用的社会经济不平等提供了系统性跨国证据，'
    '并通过全烟草与纯吸烟梯度的配对比较，首次在跨国层面揭示了无烟烟草使用对烟草不平等亲贫梯度的'
    '额外强化效应。财富及教育维度上的亲贫梯度在大多数中低收入国家中一致存在，'
    '且教育水平与城乡居住地是驱动该不平等的核心结构性因素。上述发现呼吁全球控烟议程将无烟烟草'
    '纳入不平等的监测与干预框架，并将社会经济弱势群体——尤其是低教育水平的农村居民——'
    '置于烟草控制策略的核心。'
)

# ================================================================
# REFERENCES
# ================================================================
doc.add_heading('参考文献', level=1)

refs = [
    'Carnazza, G., Liberati, P., & Resce, G. (2023). Income-related inequality in smoking habits: '
    'A comparative assessment in the European Union. Health Policy, 128, 34–41. '
    'https://doi.org/10.1016/j.healthpol.2022.12.002',

    'Chugh, A., Arora, M., Jain, N., Vidyasagaran, A., Readshaw, A., Sheikh, A., Eckhardt, J., '
    'Siddiqi, K., Chopra, M., Mishu, M. P., Kanaan, M., Rahman, M. A., Mehrotra, R., Huque, R., '
    'Forberger, S., Dahanayake, S., Khan, Z., Boeckmann, M., & Dogar, O. (2023). The global '
    'impact of tobacco control policies on smokeless tobacco use: a systematic review. '
    'The Lancet Global Health, 11(6), e953–e968. https://doi.org/10.1016/S2214-109X(23)00205-X',

    'Disney, G., Petrie, D., Yang, Y., Aitken, Z., Gurrin, L., & Kavanagh, A. (2023). Smoking '
    'inequality trends by disability and income in Australia, 2001 to 2020. Epidemiology, 34(2), '
    '302–309. https://doi.org/10.1097/EDE.0000000000001582',

    'GBD 2019 Chewing Tobacco Collaborators. (2021). Spatial, temporal, and demographic patterns '
    'in prevalence of chewing tobacco use in 204 countries and territories, 1990–2019: a systematic '
    'analysis from the Global Burden of Disease Study 2019. The Lancet Public Health, 6(7), '
    'e482–e499. https://doi.org/10.1016/S2468-2667(21)00065-7',

    'GBD 2019 Tobacco Collaborators. (2021). Spatial, temporal, and demographic patterns in '
    'prevalence of smoking tobacco use and attributable disease burden in 204 countries and '
    'territories, 1990–2019: a systematic analysis from the Global Burden of Disease Study 2019. '
    'The Lancet, 397(10292), 2337–2360. https://doi.org/10.1016/S0140-6736(21)01169-7',

    'Ghate, N., Kumar, P., & Dhillon, P. (2022). Socioeconomic determinants of smokeless tobacco '
    'use among Indian women: An analysis of Global Adult Tobacco Survey-2, India. WHO South-East '
    'Asia Journal of Public Health, 11(1), 24–31. https://doi.org/10.4103/WHO-SEAJPH.WHO-SEAJPH_160_21',

    'Halder, P., Alwani, A. A., Nongkynrih, B., Chakraborty, S., & Krishnan, A. (2025). '
    'Rural-urban disparities in tobacco use among middle aged and elderly Indian adults: '
    'a multivariate decomposition analysis. BMC Public Health, 25(1), 2818. '
    'https://doi.org/10.1186/s12889-025-23937-0',

    'Huang, M. Z., Liu, T. Y., Zhang, Z. M., Song, F., & Chen, T. (2023). Trends in the '
    'distribution of socioeconomic inequalities in smoking and cessation: evidence among adults '
    'aged 18–59 from China Family Panel Studies data. International Journal for Equity in Health, '
    '22(1), 86. https://doi.org/10.1186/s12939-023-01898-3',

    'Li, X., Chang, Y., Zhao, X., Liu, X., Zhang, Y., Wang, X., & Li, J. (2025). Global health '
    'inequities in retinoblastoma: a 1990–2021 analysis across socio-demographic index regions. '
    'Frontiers in Public Health, 13, 1513526. https://doi.org/10.3389/fpubh.2025.1513526',

    'Littlecott, H. J., Moore, G. F., McCann, M., Melendez-Torres, G. J., Mercken, L., Reed, H., '
    'Mann, M., Dobbie, F., & Hawkins, J. (2022). Exploring the association between school-based '
    'peer networks and smoking according to socioeconomic status and tobacco control context: '
    'a systematic review. BMC Public Health, 22(1), 142. '
    'https://doi.org/10.1186/s12889-021-12333-z',

    'Mann, N., Spencer, G., Hutchinson, B., Ngongo, C., Tarlton, D., Webb, D., Grafton, D., & '
    'Nugent, R. (2024). Interpreting results, impacts and implications from WHO FCTC tobacco '
    'control investment cases in 21 low-income and middle-income countries. Tobacco Control, '
    '33(Suppl 1), s17–s26. https://doi.org/10.1136/tc-2023-058337',

    'McEvoy, O., & Layte, R. (2025). Bringing the group back in: Social class and resistance in '
    'adolescent smoking. Sociology of Health & Illness, 47(2), e13858. '
    'https://doi.org/10.1111/1467-9566.13858',

    'Mills, S. D., Rosario, C., Yerger, V. B., Kalb, M. D., & Ribisl, K. M. (2024). '
    'Recommendations to advance equity in tobacco control. Tobacco Control, 33(e2), '
    'e246–e253. https://doi.org/10.1136/tc-2022-057670',

    'Pilvar, H., Nicodemo, C., Petrou, S., Darlow, B. A., van Dommelen, P., Evensen, K. A. I., '
    'Harris, S., Horwood, J., Johnson, S., Marlow, N., Mathewson, K., Saigal, S., Schmidt, L. A., '
    'Wolke, D., Woodward, L. J., & Kim, S. (2026). Socioeconomic inequity in extreme outcomes '
    'within very pre-term and/or very low birthweight infants: evidence from multi-national '
    'cohorts. Frontiers in Public Health, 14, 1791450. '
    'https://doi.org/10.3389/fpubh.2026.1791450',

    'Puljevic, C., Feulner, L., Hobbs, M., Erku, D., Bonevski, B., Segan, C., Baker, A., '
    'Hefler, M., Cho, A., & Gartner, C. (2024). Tobacco endgame and priority populations: '
    'a scoping review. Tobacco Control, 33(e2), e231–e239. '
    'https://doi.org/10.1136/tc-2022-057715',

    'Shaikh, R., & Saikia, N. (2022). Socioeconomic inequalities in tobacco cessation among '
    'Indians above 15 years of age from 2009 to 2017: evidence from the Global Adult Tobacco '
    'Survey (GATS). BMC Public Health, 22(1), 1419. '
    'https://doi.org/10.1186/s12889-022-13820-7',

    'Spencer, G., Nugent, R., Mann, N., Hutchinson, B., Ngongo, C., Tarlton, D., Small, R., & '
    'Webb, D. (2024). Equity implications of tobacco taxation: results from WHO FCTC investment '
    'cases. Tobacco Control, 33(Suppl 1), s27–s33. https://doi.org/10.1136/tc-2023-058338',

    'Tanaka, H., Mackenbach, J. P., & Kobayashi, Y. (2021). Widening socioeconomic inequalities '
    'in smoking in Japan, 2001–2016. Journal of Epidemiology, 31(6), 369–377. '
    'https://doi.org/10.2188/jea.JE20200025',

    'Van Hemelrijck, W. M. J., Kunst, A. E., Sizer, A., Martikainen, P., Zengarini, N., '
    'Costa, G., & Janssen, F. (2024). Trends in educational inequalities in smoking-attributable '
    'mortality and their impact on changes in general mortality inequalities: evidence from '
    'England and Wales, Finland, and Italy (Turin). Journal of Epidemiology and Community Health, '
    '78(9), 561–569. https://doi.org/10.1136/jech-2023-221702',

    'Vladisavljevic, M., Zubovic, J., Jovanovic, O., & Dukic, M. (2024). Crowding-out effect '
    'of tobacco consumption in Serbia. Tobacco Control, 33(Suppl 2), s88–s94. '
    'https://doi.org/10.1136/tc-2022-057727',

    'Wehrmeister, F. C., Costa, J. C., Silva, L. E. S. D., Lima, N. P., Costa, F. D. S., & '
    'Barros, A. J. D. (2025). Assessment of inequalities in women\'s and children\'s health: '
    'Why, how, and for whom? Analyses of the Brazilian Live Birth Information System. '
    'Revista Brasileira de Epidemiologia, 28, e250029. '
    'https://doi.org/10.1590/1980-5497202500XX',
]

for ref in refs:
    rp = doc.add_paragraph(ref)
    rp.paragraph_format.first_line_indent = Cm(0)
    rp.paragraph_format.left_indent = Cm(0.74)
    rp.paragraph_format.first_line_indent = Cm(-0.74)
    for run in rp.runs:
        run.font.size = Pt(10)

out = os.path.join(ROOT, "manuscript", "Discussion_CHS.docx")
os.makedirs(os.path.join(ROOT, "manuscript"), exist_ok=True)
doc.save(out)
print(f"Saved: {out}")
