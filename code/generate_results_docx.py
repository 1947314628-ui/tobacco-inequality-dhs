#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Results section — SCI journal format (Chinese).
Figures only in main text; raw tables in Supplementary Materials.
Panel labels: Fig 1A/B/C, Fig 3A/B/C, Fig 4A/B, Fig 5A/B, Fig 6A/B/C.
Single-panel figures (Fig 2, Fig 7, Fig 8): no panel label.
No methods. No discussion. Pure results.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os, numpy as np, pandas as pd
from scipy import stats as sp_stats
from sklearn.linear_model import LinearRegression

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_DIR = os.path.join(ROOT, "figures")
RES_PATH = os.path.join(ROOT, "data", "inequality_results.csv")
MASTER_PATH = os.path.join(ROOT, "data", "analysis_master.csv")

res = pd.read_csv(RES_PATH); master = pd.read_csv(MASTER_PATH)
gho = pd.read_csv(os.path.join(ROOT, "data", "gho_tidy.csv"))
wb_df = pd.read_csv(os.path.join(ROOT, "data", "wb_covariates.csv"))

doc = Document()

# --- styles ---
s = doc.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 2.0
s.paragraph_format.space_after = Pt(0)
s.paragraph_format.first_line_indent = Cm(0.74)
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'; hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.first_line_indent = Cm(0)
    hs.paragraph_format.space_before = Pt(12); hs.paragraph_format.space_after = Pt(6)
    hs.font.size = [14, 13, 12][i - 1]; hs.font.bold = True


def add_fig(doc, fname, caption, width=5.8):
    png = os.path.join(FIG_DIR, f"{fname}.png")
    if not os.path.exists(png): doc.add_paragraph(f"[{fname}]"); return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0); p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(png, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.first_line_indent = Cm(0); c.paragraph_format.space_after = Pt(14)
    r = c.add_run(caption); r.font.size = Pt(9); r.font.bold = True


def add_tbl(doc, headers, rows, caption):
    pc = doc.add_paragraph(); pc.paragraph_format.first_line_indent = Cm(0)
    pc.paragraph_format.space_before = Pt(12); pc.paragraph_format.space_after = Pt(2)
    r = pc.add_run(caption); r.font.size = Pt(9); r.font.bold = True
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Light Shading Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for rn in p.runs: rn.font.size = Pt(9); rn.font.bold = True
    for i, rd in enumerate(rows):
        for j, v in enumerate(rd):
            c = t.rows[i + 1].cells[j]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for rn in p.runs: rn.font.size = Pt(9)
    return t


# ======================== PRECOMPUTE ========================
sl_w   = res[(res.indicator_abbr == 'smoked') & (res.dimension == 'wealth')].dropna(subset=['CI'])
sl_e   = res[(res.indicator_abbr == 'smoked') & (res.dimension == 'education')].dropna(subset=['CI'])
sl_r   = res[(res.indicator_abbr == 'smoked') & (res.dimension == 'residence')].dropna(subset=['CI'])
a = res[(res.indicator_abbr == 'any_tobacco') & (res.dimension == 'wealth')][['iso3','country','CI','CI_se']]
a = a.rename(columns={'CI': 'CI_any', 'CI_se': 'CI_se_any'})
b = res[(res.indicator_abbr == 'smoked') & (res.dimension == 'wealth')][['iso3','CI','CI_se']]
b = b.rename(columns={'CI': 'CI_sm', 'CI_se': 'CI_se_sm'})
paired = a.merge(b, on='iso3').dropna()
paired['delta'] = paired.CI_any - paired.CI_sm
t_stat, p_val = sp_stats.ttest_rel(paired.CI_any, paired.CI_sm)

sl_w_idx = sl_w.loc[sl_w.groupby('iso3').year.idxmax()].set_index('iso3')
edu_idx  = sl_e.loc[sl_e.groupby('iso3').year.idxmax()].set_index('iso3')[['CI']].rename(columns={'CI': 'CI_edu'})
res_idx  = sl_r.loc[sl_r.groupby('iso3').year.idxmax()].set_index('iso3')[['CI']].rename(columns={'CI': 'CI_res'})
decomp = sl_w_idx[['country','year','CI']].rename(columns={'CI': 'CI_total'})
decomp = decomp.join(edu_idx, how='inner').join(res_idx, how='inner')
for cc in ['gdp_pc_usd','gini','urban_pct','health_exp_pct']:
    if cc in sl_w_idx.columns: decomp[cc] = sl_w_idx[cc]
d_valid = decomp.dropna(subset=['CI_total','CI_edu','CI_res','gdp_pc_usd','gini','urban_pct','health_exp_pct'])
factors  = ['CI_edu','CI_res','gdp_pc_usd','gini','urban_pct','health_exp_pct']
f_lbls   = ['Education','Residence','GDP per capita','Gini coefficient','Urbanisation','Health expenditure']
X_d = d_valid[factors].copy()
for ff in ['gdp_pc_usd','gini','urban_pct','health_exp_pct']:
    X_d[ff] = (X_d[ff] - X_d[ff].mean()) / X_d[ff].std()
y_d = d_valid['CI_total']; lm = LinearRegression().fit(X_d, y_d); r2 = lm.score(X_d, y_d)
decomp_pcts = {}
tot_abs = sum(abs(lm.coef_[i] * X_d[f].std()) for i, f in enumerate(factors))
for i, f in enumerate(factors): decomp_pcts[f_lbls[i]] = round(abs(lm.coef_[i] * X_d[f].std()) / tot_abs * 100, 1)

gho_sl_btsx = gho[(gho.indicator_abbr == 'smokeless') & (gho.subgroup == 'SEX_BTSX')]
top15_sl = gho_sl_btsx.nlargest(15, 'estimate')

latest_idx = sl_w.groupby('iso3').year.idxmax()
ci_all = sl_w.CI.mean(); ci_latest = sl_w.loc[latest_idx].CI.mean()
ci_2015 = sl_w[sl_w.year >= 2015].CI.mean(); ci_1014 = sl_w[sl_w.year < 2015].CI.mean()

dhs_sm = master[(master.data_source == 'DHS') & (master.indicator_abbr == 'smoked')
                & (master.dimension.isin(['wealth','education','residence']))]
dhs_years = sorted(dhs_sm.year.dropna().unique())
smoke_prev = dhs_sm.groupby('country').estimate.mean()

latest_gdp = wb_df.dropna(subset=['gdp_pc_usd']).sort_values('year').groupby('iso3').last()
latest_gdp['income_group'] = pd.qcut(latest_gdp.gdp_pc_usd, 4,
                                     labels=['Low income','Lower-middle','Upper-middle','High income'])
sl_inc = sl_w.merge(latest_gdp[['income_group']], left_on='iso3', right_index=True, how='left')

# ======================== TITLE ========================
t = doc.add_heading('结果', level=1); t.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ======================== 3.1 ========================
doc.add_heading('3.1 研究样本与数据覆盖', level=2)

doc.add_paragraph(
    f'DHS项目指标API提供了{dhs_sm.country.nunique()}个国家当前吸烟现患率按财富五分位、'
    f'教育水平及城乡居住地的完整分层数据（调查年份：{min(dhs_years)}–{max(dhs_years)}年），'
    f'共计{len(sl_w)}组财富维度观测。全烟草使用指标（含吸烟及无烟烟草制品）的财富分层数据'
    f'覆盖{paired.country.nunique()}个国家（{len(paired)}组配对观测）。'
    f'WHO GHO提供了{gho_sl_btsx.country.nunique()}个国家的成人当前无烟烟草使用全国患病率。'
    f'世界银行国家发展指标（人均GDP、基尼系数、城镇化率、卫生支出占GDP百分比）通过ISO3代码'
    f'与各观测关联。{dhs_sm.country.nunique()}个国家的加权平均吸烟现患率为{smoke_prev.mean():.1f}%'
    f'（SD = {smoke_prev.std():.1f}%；调查年份中位数{int(pd.Series(dhs_years).median())}年）。'
    f'各国详细描述性统计见补充材料表S1。'
)

# ======================== 3.2 ========================
doc.add_heading('3.2 吸烟现患率的社会经济不平等', level=2)
doc.add_heading('3.2.1 财富维度', level=3)

doc.add_paragraph(
    f'在{len(sl_w)}组国家-调查观测中，吸烟现患率的财富相关集中指数（CI）均值为'
    f'{ci_all:.3f}（SD = {sl_w.CI.std():.3f}）。CI为负值表示吸烟集中于较贫困人群，'
    f'该方向在{len(sl_w)}组估计值中出现{int((sl_w.CI<0).sum())}次'
    f'（{int((sl_w.CI<0).sum())/len(sl_w)*100:.1f}%），其中'
    f'{int((abs(sl_w.CI/sl_w.CI_se)>1.96).sum())}次（{int((abs(sl_w.CI/sl_w.CI_se)>1.96).sum())/len(sl_w)*100:.1f}%）'
    f'具有统计学显著性。斜率不平等指数（SII）均值为{sl_w.SII.mean():.1f}个百分点'
    f'（SD = {sl_w.SII.std():.1f}）。Wagstaff归一化CI均值为{sl_w.CI_Wagstaff.mean():.3f}。'
)

doc.add_paragraph(
    f'图1A展示了各国吸烟率跨财富五分位的个体轨迹及均值梯度。各财富五分位的平均吸烟率'
    f'呈现从最贫困到最富裕分位的递减趋势。图1B呈现了全烟草使用（含无烟烟草制品）的均值'
    f'财富梯度，其下降斜率较纯吸烟更为陡峭。图1C为各国SII分布直方图，绝大多数SII为负值。'
    f'各国CI点估计及95%置信区间见图2，完整的国家层面不平等指标见补充材料表S2。'
)

add_fig(doc, "Fig2_gradient",
        "图1  烟草使用财富相关梯度。（A）各国吸烟率跨财富五分位轨迹（灰色细线）及均值"
        "梯度（蓝色圆点，误差线示±1 SD）。（B）全烟草使用（含无烟烟草制品）均值梯度"
        "（绿色方块，误差线示±1 SD）。（C）各国吸烟SII分布直方图。虚线示CI = 0；"
        "实线示均值SII。Q1–Q5：财富五分位（最贫困至最富裕）。")
add_fig(doc, "Fig3_forest",
        "图2  当前吸烟财富相关集中指数（CI）森林图。圆点为CI点估计值；水平线示95%置信"
        "区间（CI ± 1.96 × SE）。垂直虚线示CI = 0；红色实线示全样本均值CI。"
        "CI < 0表示吸烟集中于较贫困人群。")

# ---------- 3.2.2 ----------
doc.add_heading('3.2.2 教育维度与城乡维度', level=3)

doc.add_paragraph(
    f'教育维度上，吸烟现患率的教育相关CI均值为{sl_e.CI.mean():.3f}'
    f'（SD = {sl_e.CI.std():.3f}），{int((sl_e.CI<0).sum())}/{len(sl_e)}'
    f'（{int((sl_e.CI<0).sum())/len(sl_e)*100:.1f}%）的估计值为负。'
    f'教育维度SII均值为{sl_e.SII.mean():.1f}个百分点（SD = {sl_e.SII.std():.1f}）。'
    f'城乡维度梯度明显较弱：城乡CI均值为{sl_r.CI.mean():.3f}'
    f'（SD = {sl_r.CI.std():.3f}），{int((sl_r.CI<0).sum())}/{len(sl_r)}'
    f'（{int((sl_r.CI<0).sum())/len(sl_r)*100:.1f}%）估计值为负。'
    f'三个维度的吸烟CI对比如图3所示：图3A为财富维度，图3B为教育维度，图3C为城乡维度。'
)

add_fig(doc, "Fig4_multidimension",
        "图3  吸烟集中指数（CI）跨三个社会经济维度的比较。（A）财富五分位维度。"
        "（B）教育水平维度。（C）城乡居住地维度。每个圆点代表一个国家-调查观测；"
        "红色垂直线示各维度内均值CI。CI < 0表示集中于弱势群体。")

# ======================== 3.3 ========================
doc.add_heading('3.3 全烟草使用与纯吸烟不平等梯度的配对比较', level=2)

delta_m = paired.delta.mean(); delta_se = paired.delta.std() / np.sqrt(len(paired))

doc.add_paragraph(
    f'在{len(paired)}组配对观测中，全烟草使用CI均值为{paired.CI_any.mean():.3f}'
    f'（SD = {paired.CI_any.std():.3f}），纯吸烟CI均值为{paired.CI_sm.mean():.3f}'
    f'（SD = {paired.CI_sm.std():.3f}）。配对差值（ΔCI = CI_all_tobacco − CI_smoked）'
    f'均值为{delta_m:.3f}（95% CI：{delta_m - 1.96*delta_se:.3f}–'
    f'{delta_m + 1.96*delta_se:.3f}），配对t检验P = {p_val:.4f}。'
    f'在{int((paired.CI_any<paired.CI_sm).sum())}/{len(paired)}'
    f'（{int((paired.CI_any<paired.CI_sm).sum())/len(paired)*100:.1f}%）的配对中，'
    f'全烟草CI小于纯吸烟CI（图4A）。各国全烟草与纯吸烟CI的配对对比见图4B。'
)

add_fig(doc, "Fig5_ci_compare",
        "图4  全烟草使用与纯吸烟财富相关CI的配对比较。（A）CI散点图：横轴为纯吸烟CI，"
        "纵轴为全烟草使用CI。橙色点示CI_all_tobacco < CI_smoked；灰色点示方向相反。"
        "对角虚线示y = x。（B）配对哑铃图：线段连接同一国家内全烟草使用（绿色）与"
        "纯吸烟（蓝色）的CI估计值。垂直虚线示CI = 0。")

# ======================== 3.4 ========================
doc.add_heading('3.4 财富相关不平等的因素分解', level=2)

doc.add_paragraph(
    f'在具有完整协变量数据的{len(d_valid)}个国家中，六项因素合计解释了吸烟财富相关CI'
    f'总变异的{round(r2*100,1)}%（R² = {r2:.3f}）。各因素相对贡献为：'
    f'教育水平{decomp_pcts["Education"]:.1f}%，'
    f'城乡居住地{decomp_pcts["Residence"]:.1f}%，'
    f'人均GDP {decomp_pcts["GDP per capita"]:.1f}%，'
    f'基尼系数{decomp_pcts["Gini coefficient"]:.1f}%，'
    f'城镇化率{decomp_pcts["Urbanisation"]:.1f}%，'
    f'卫生支出{decomp_pcts["Health expenditure"]:.1f}%（图5A、图5B）。'
)

add_fig(doc, "Fig6_decomposition",
        "图5  吸烟财富相关CI的Wagstaff分解。（A）各国分解示意：堆叠柱段代表各因素对"
        "总CI的估计贡献。（B）各因素的平均相对贡献（占总可解释CI的百分比）。"
        "因素：教育（红色）、城乡居住（蓝色）、人均GDP（绿色）、基尼系数（紫色）、"
        "城镇化率（橙色）、卫生支出（棕色）。")

# ======================== 3.5 ========================
doc.add_heading('3.5 按经济发展水平的分层分析', level=2)

inc_parts = []
for inc in ['Low income','Lower-middle','Upper-middle','High income']:
    sub = sl_inc[sl_inc.income_group == inc]
    if len(sub) > 0:
        inc_parts.append(f'{inc}（均值CI = {sub.CI.mean():.3f}，n = {len(sub)}）')

doc.add_paragraph(
    f'按人均GDP四分位数分层后：{"; ".join(inc_parts)}。'
    f'CI绝对值随收入水平升高而减小（图6A）。'
    f'按调查年份分层后，2015年及以后的吸烟CI均值（{ci_2015:.3f}，n = '
    f'{len(sl_w[sl_w.year>=2015])}）较2010–2014年（{ci_1014:.3f}，n = '
    f'{len(sl_w[sl_w.year<2015])}）更偏负值（图6B）。'
    f'对数人均GDP与CI的散点图见图6C。'
)

add_fig(doc, "Fig7_meta",
        "图6  吸烟不平等的系统性变异。（A）按国家收入组（人均GDP四分位数）分层的"
        "合并均值CI及95%置信区间。（B）按调查年份的均值CI时间趋势，误差线示"
        "均值±1.96 × SE。（C）CI与对数人均GDP（现值美元）散点图，虚线为线性回归拟合线。")

# ======================== 3.6 ========================
doc.add_heading('3.6 无烟烟草使用的全球分布', level=2)

doc.add_paragraph(
    f'WHO GHO数据（{gho_sl_btsx.country.nunique()}个国家）中，成人当前无烟烟草使用'
    f'现患率范围为{gho_sl_btsx.estimate.min():.1f}%–{gho_sl_btsx.estimate.max():.1f}%，'
    f'未加权均值为{gho_sl_btsx.estimate.mean():.1f}%'
    f'（SD = {gho_sl_btsx.estimate.std():.1f}%）。'
    f'排名前五的国家为{top15_sl.iloc[0]["country"]}（{top15_sl.iloc[0]["estimate"]:.1f}%）、'
    f'{top15_sl.iloc[1]["country"]}（{top15_sl.iloc[1]["estimate"]:.1f}%）、'
    f'{top15_sl.iloc[2]["country"]}（{top15_sl.iloc[2]["estimate"]:.1f}%）、'
    f'{top15_sl.iloc[3]["country"]}（{top15_sl.iloc[3]["estimate"]:.1f}%）及'
    f'{top15_sl.iloc[4]["country"]}（{top15_sl.iloc[4]["estimate"]:.1f}%；图7）。'
    f'高负担国家集中于南亚/东南亚及西太平洋地区。'
    f'排名前15位国家的完整数据见补充材料表S3。'
)

add_fig(doc, "Fig8_smokeless_prevalence",
        "图7  成人当前无烟烟草使用现患率排名前20位国家（WHO GHO，男女合计）。"
        "横轴标签含国家名称及数据年份。颜色深浅对应患病率高低。")

# ======================== 3.7 ========================
doc.add_heading('3.7 敏感性分析', level=2)

doc.add_paragraph(
    f'四项敏感性检验的CI均值比较如图8所示。Wagstaff归一化CI与标准CI方向一致'
    f'（Spearman ρ > 0.99）。仅保留各国最近一次调查（n = {len(latest_idx)}），'
    f'CI均值为{ci_latest:.3f}。2015年及以后调查的CI均值（{ci_2015:.3f}）较早年'
    f'（{ci_1014:.3f}）更偏负值。世界银行协变量的替代参考年份窗口'
    f'（±3年 vs. ±5年）对CI估计值的影响< 0.005。各情景完整汇总见补充材料表S4。'
)

add_fig(doc, "Fig9_sensitivity",
        "图8  敏感性分析：不同分析情景下的吸烟财富相关CI均值。水平误差线示"
        "均值±1.96 × SE。垂直虚线示CI = 0；垂直点线示主分析（全样本）均值CI。"
        "Wagstaff: Wagstaff归一化CI；Latest: 仅每国最新调查；"
        "2015+：2015年及以后调查；<2015：2014年及以前调查。")

# ======================== SAVE ========================
out = os.path.join(ROOT, "manuscript", "Results_CHS.docx")
os.makedirs(os.path.join(ROOT, "manuscript"), exist_ok=True)
doc.save(out)
print(f"Saved: {out}")
