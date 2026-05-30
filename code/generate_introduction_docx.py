#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate Introduction section — ALL citations verified from reference files.
APA format. No discussion/methods. Past tense only.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
doc = Document()

# --- Styles ---
s = doc.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 2.0
s.paragraph_format.space_after = Pt(0)
s.paragraph_format.first_line_indent = Cm(0.74)

for i in range(1, 3):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'; hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.first_line_indent = Cm(0)
    hs.paragraph_format.space_before = Pt(10); hs.paragraph_format.space_after = Pt(4)
    hs.font.size = [13, 12][i - 1]; hs.font.bold = True


def p(text):
    doc.add_paragraph(text)


# ================================================================
t = doc.add_heading('引言', level=1)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ================================================================
# Para 1: Oral cancer + SLT burden
# ================================================================
p(
    '唇及口腔癌是全球癌症负担的重要组成部分。据全球疾病负担研究（GBD）2019年数据，'
    '2019年全球新发唇、口腔及咽部癌症病例约37万例，年龄标准化发病率为每10万人4.8例，'
    '1990年至2019年间发病例数增长了约120%，其中超过90%的疾病负担集中于南亚及东南亚地区'
    '（GBD 2019 Lip, Oral, and Pharyngeal Cancer Collaborators, 2023）。'
    'Rumgay等（2024）基于GLOBOCAN 2022数据的系统分析显示，全球约30.8%的口腔癌新发病例'
    '（约120,200例）可归因于无烟烟草及槟榔坚果的使用，且该归因比例在东南亚区域高达50%以上，'
    '其中90.2%的归因病例集中于中低收入国家。上述证据表明，无烟烟草使用构成了口腔癌'
    '全球疾病负担中一个高度可预防、但尚未得到充分关注的关键风险因素。'
)

# ================================================================
# Para 2: Smoking inequality gradient — well established
# ================================================================
p(
    '烟草使用的社会经济不平等是公共卫生研究的经典议题。大量基于个体层面调查数据的跨国研究'
    '一致表明，吸烟现患率在多数中低收入国家呈现亲贫（pro-poor）梯度，'
    '即较低社会经济地位群体的吸烟率显著高于较高社会经济地位群体'
    '（GBD 2019 Tobacco Collaborators, 2021; Sreeramareddy & Acharya, 2021）。'
    'Sreeramareddy与Acharya（2021）利用DHS数据对22个撒哈拉以南非洲国家的分析发现，'
    '以家庭财富指数及教育水平为不平等标志时，男性吸烟率的绝对社会经济不平等幅度约为女性的3倍，'
    '且低教育水平及低收入群体的烟草使用率显著更高。Dai等（2022）对过去半个世纪全球吸烟流行'
    '演变的分析进一步指出，吸烟率下降在高收入国家已取得显著进展，但在中低收入国家的进展缓慢，'
    '导致烟草相关疾病负担的不平等在全球范围内持续存在并可能扩大。'
    'Chen等（2021）基于19个中低收入国家DHS数据的分析发现，低教育水平及低收入状态与男性'
    '多重烟草使用（dual and poly-tobacco use）显著相关，提示社会经济弱势群体不仅吸烟率更高，'
    '还可能同时使用多种烟草制品。'
)

# ================================================================
# Para 3: SLT may have a DIFFERENT, steeper gradient
# ================================================================
p(
    '然而，无烟烟草使用的社会经济分布模式可能与吸烟存在本质差异。从制品特征来看，'
    '无烟烟草制品——尤其是南亚地区广泛消费的低价、小包装gutka及槟榔嚼块——在街边零售点'
    '高度可及，不受到与吸烟同等的公共场所禁令限制，且往往被视为文化习俗或社交传统的一部分'
    '（Kaur et al., 2024）。Yang等（2022）基于全球青少年烟草调查（GYTS）数据，'
    '在138个国家/地区12–16岁青少年中报告了4.4%的无烟烟草使用率，东南亚区域最高达6.1%。'
    '从经济机制来看，Spencer等（2024）基于WHO FCTC投资案例对19个国家的分析发现，'
    '烟草税的价格弹性在低收入群体中最大——30%的价格提升可使最贫困20%人口的吸烟率下降幅度'
    '最大，而该群体仅承担12%的新增税收支出，表明通过价格杠杆的控烟收益具有亲贫属性。'
    'Vladisavljevic等（2024）及Macias Sanchez与Garcia Gomez（2024）分别在塞尔维亚及墨西哥'
    '利用住户预算调查数据证实了烟草消费对低收入家庭的挤出效应（crowding-out effect）——'
    '烟草支出显著挤占了食品、衣物、教育及医疗等基本需求，且这种挤出效应在低收入家庭中'
    '更为显著，部分家庭因烟草支出而落入贫困线以下。上述证据共同提示，无烟烟草使用的社会经济'
    '梯度可能比吸烟更为陡峭——即更强烈地集中于社会经济弱势群体。然而，该假设尚未在跨国可比'
    '数据框架下得到直接的量化检验。'
)

# ================================================================
# Para 4: Research gap
# ================================================================
p(
    '综上所述，现有文献在以下方面存在明显空白。第一，尽管已有研究分别探讨了吸烟的社会经济'
    '不平等（Sreeramareddy & Acharya, 2021; Chen et al., 2021）或无烟烟草使用的流行病学'
    '分布（Rumgay et al., 2024; Kaur et al., 2024），但缺乏采用统一的不平等测量框架'
    '（如集中指数、斜率不平等指数）对吸烟与无烟烟草使用进行直接跨国比较的研究。'
    '第二，驱动烟草使用财富相关不平等的结构性因素——包括教育水平、城乡居住地及国家宏观'
    '经济发展指标——各自的相对贡献尚未在跨国背景下被量化分解。第三，不平等程度是否随国家收入'
    '水平、WHO区域或调查时期呈现系统性变异仍不清楚。上述空白的填补对于在全球控烟议程中'
    '制定差异化干预策略——尤其是精准靶向社会经济弱势群体的无烟烟草控制措施——具有直接的政策意义'
    '（Spencer et al., 2024; Kaur et al., 2024）。'
)

# ================================================================
# Para 5: Aims
# ================================================================
p(
    '因此，本研究旨在利用可公开获取的跨国聚合数据，系统量化无烟烟草使用及吸烟在财富、'
    '教育及城乡维度上的社会经济不平等程度。具体目标为：(1) 估算各国家的绝对不平等指标'
    '（斜率不平等指数，SII）及相对不平等指标（相对不平等指数RII、集中指数CI及Wagstaff归一化CI），'
    '涵盖多个社会经济维度；(2) 直接检验全烟草使用（含无烟烟草制品）的社会经济梯度是否'
    '比纯吸烟更为亲贫（即CI_all_tobacco < CI_smoked）；(3) 采用Wagstaff分解方法量化'
    '各社会经济决定因素对财富相关不平等的相对贡献；(4) 按国家收入水平及调查时期分层，'
    '考察不平等指数的系统性变异。'
)

# ================================================================
# REFERENCES (APA format)
# ================================================================
doc.add_heading('参考文献', level=1)

refs = [
    'Chen, D. T., Millett, C., & Filippidis, F. T. (2021). Prevalence and determinants of dual '
    'and poly-tobacco use among males in 19 low-and middle-income countries: Implications for a '
    'comprehensive tobacco control regulation. Preventive Medicine, 142, 106377. '
    'https://doi.org/10.1016/j.ypmed.2020.106377',

    'Dai, X., Gakidou, E., & Lopez, A. D. (2022). Evolution of the global smoking epidemic over '
    'the past half century: Strengthening the evidence base for policy action. Tobacco Control, '
    '31(2), 129–137. https://doi.org/10.1136/tobaccocontrol-2021-056535',

    'GBD 2019 Lip, Oral, and Pharyngeal Cancer Collaborators. (2023). The global, regional, and '
    'national burden of adult lip, oral, and pharyngeal cancer in 204 countries and territories: '
    'A systematic analysis for the Global Burden of Disease Study 2019. JAMA Oncology, 9(10), '
    '1401–1416. https://doi.org/10.1001/jamaoncol.2023.2960',

    'GBD 2019 Tobacco Collaborators. (2021). Spatial, temporal, and demographic patterns in '
    'prevalence of smoking tobacco use and attributable disease burden in 204 countries and '
    'territories, 1990–2019: A systematic analysis from the Global Burden of Disease Study 2019. '
    'The Lancet, 397(10292), 2337–2360. https://doi.org/10.1016/S0140-6736(21)01169-7',

    'Kaur, J., Rinkoo, A. V., & Richardson, S. (2024). Trends in smokeless tobacco use and '
    'attributable mortality and morbidity in the South-East Asia Region: Implications for policy. '
    'Tobacco Control, 33(4), 425–433. https://doi.org/10.1136/tc-2022-057669',

    'Macias Sanchez, A., & Garcia Gomez, A. (2024). Crowding out and impoverishing effect of '
    'tobacco in Mexico. Tobacco Control, 33(Suppl 2), s68–s74. '
    'https://doi.org/10.1136/tc-2022-057791',

    'Moghadam, T. Z., Zandian, H., Fazlzadeh, M., Kalan, M. E., & Pourfarzi, F. (2023). '
    'Socioeconomic and environmental factors associated with waterpipe tobacco smoking among '
    'Iranian adults: A PERSIAN cohort-based cross-sectional study. BMC Public Health, 23(1), '
    '1295. https://doi.org/10.1186/s12889-023-16176-8',

    'Rumgay, H., Nethan, S. T., Shah, R., Vignat, J., Ayo-Yusuf, O., Chaturvedi, P., '
    'Guerra, E. N. S., Gupta, P. C., Gupta, R., Liu, S., Magnusson, C., Parascandola, M., '
    'Paulino, Y. C., Rezaei, N., Siddiqi, K., Warnakulasuriya, S., Lauby-Secretan, B., & '
    'Soerjomataram, I. (2024). Global burden of oral cancer in 2022 attributable to smokeless '
    'tobacco and areca nut consumption: A population attributable fraction analysis. '
    'The Lancet Oncology, 25(11), 1413–1423. https://doi.org/10.1016/S1470-2045(24)00458-3',

    'Spencer, G., Nugent, R., Mann, N., Hutchinson, B., Ngongo, C., Tarlton, D., Small, R., & '
    'Webb, D. (2024). Equity implications of tobacco taxation: Results from WHO FCTC investment '
    'cases. Tobacco Control, 33(Suppl 1), s27–s33. https://doi.org/10.1136/tc-2023-058338',

    'Sreeramareddy, C. T., & Acharya, K. (2021). Trends in prevalence of tobacco use by sex and '
    'socioeconomic status in 22 Sub-Saharan African countries, 2003–2019. JAMA Network Open, '
    '4(12), e2137820. https://doi.org/10.1001/jamanetworkopen.2021.37820',

    'Vladisavljevic, M., Zubovic, J., Jovanovic, O., & Dukic, M. (2024). Crowding-out effect '
    'of tobacco consumption in Serbia. Tobacco Control, 33(Suppl 2), s88–s94. '
    'https://doi.org/10.1136/tc-2022-057727',

    'Yang, H., Ma, C., Zhao, M., Magnussen, C. G., & Xi, B. (2022). Prevalence and trend of '
    'smokeless tobacco use and its associated factors among adolescents aged 12–16 years in 138 '
    'countries/territories, 1999–2019. BMC Medicine, 20(1), 460. '
    'https://doi.org/10.1186/s12916-022-02662-0',

    'Zhao, Y., & Fei, L. (2025). Smoking-attributable neurological health loss: Age-specific '
    'burden and health disparities. Journal of Neurology, Neurosurgery & Psychiatry, 96(10), '
    '937–946. https://doi.org/10.1136/jnnp-2024-335536',
]

for ref in refs:
    rp = doc.add_paragraph(ref)
    rp.paragraph_format.first_line_indent = Cm(0)
    rp.paragraph_format.left_indent = Cm(0.74)
    # Hanging indent via negative first-line
    rp.paragraph_format.first_line_indent = Cm(-0.74)
    for run in rp.runs:
        run.font.size = Pt(10)

# ================================================================
out = os.path.join(ROOT, "manuscript", "Introduction_CHS.docx")
os.makedirs(os.path.join(ROOT, "manuscript"), exist_ok=True)
doc.save(out)
print(f"Saved: {out}")
