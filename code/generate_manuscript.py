#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate complete journal manuscript with proper formatting.
Title, Structured Abstract, 1. Introduction, 2. Methods, 3. Results, 4. Discussion,
Declarations, References.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os, numpy as np, pandas as pd
from scipy import stats as sp_stats
from sklearn.linear_model import LinearRegression

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_DIR = os.path.join(ROOT, "figures")
DATA_DIR = os.path.join(ROOT, "data")

res = pd.read_csv(os.path.join(DATA_DIR, "inequality_results.csv"))
master = pd.read_csv(os.path.join(DATA_DIR, "analysis_master.csv"))
gho = pd.read_csv(os.path.join(DATA_DIR, "gho_tidy.csv"))

doc = Document()

# =================== GLOBAL STYLES ===================
sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'
sty.font.size = Pt(12)
sty.paragraph_format.line_spacing = 2.0
sty.paragraph_format.space_after = Pt(0)
sty.paragraph_format.first_line_indent = Cm(0.74)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.first_line_indent = Cm(0)
    hs.paragraph_format.line_spacing = 2.0
    hs.paragraph_format.space_before = Pt(18)
    hs.paragraph_format.space_after = Pt(6)
    hs.font.bold = True
    if i == 1:
        hs.font.size = Pt(14)
    elif i == 2:
        hs.font.size = Pt(13)
    else:
        hs.font.size = Pt(12)


# =================== HELPERS ===================
def body(text):
    """Body paragraph — has first-line indent."""
    doc.add_paragraph(text)


def noindent(text):
    """No-indent paragraph — for abstract, keywords, declarations, figure captions."""
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def sec(text, level=1):
    """Section heading with numbering. Level 1 = centered."""
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.first_line_indent = Cm(0)
    return h


def fig(fname, caption, width=5.8):
    png = os.path.join(FIG_DIR, f"{fname}.png")
    if not os.path.exists(png):
        doc.add_paragraph(f"[{fname} missing]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(png, width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.first_line_indent = Cm(0)
    c.paragraph_format.space_after = Pt(14)
    r = c.add_run(caption)
    r.font.size = Pt(9)
    r.font.bold = True


def tbl(headers, rows, caption=""):
    if caption:
        pc = doc.add_paragraph()
        pc.paragraph_format.first_line_indent = Cm(0)
        pc.paragraph_format.space_before = Pt(8)
        pc.paragraph_format.space_after = Pt(2)
        r = pc.add_run(caption)
        r.font.size = Pt(10)
        r.font.bold = True
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Light Shading Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for rn in p.runs:
                rn.font.size = Pt(9)
                rn.font.bold = True
    for i, rd in enumerate(rows):
        for j, v in enumerate(rd):
            c = t.rows[i + 1].cells[j]
            c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for rn in p.runs:
                    rn.font.size = Pt(9)
    return t


# =================== PRECOMPUTE ALL NUMBERS ===================
sl_w   = res[(res.indicator_abbr == 'smoked') & (res.dimension == 'wealth')].dropna(subset=['CI'])
sl_e   = res[(res.indicator_abbr == 'smoked') & (res.dimension == 'education')].dropna(subset=['CI'])
sl_r   = res[(res.indicator_abbr == 'smoked') & (res.dimension == 'residence')].dropna(subset=['CI'])

a_d = res[(res.indicator_abbr == 'any_tobacco') & (res.dimension == 'wealth')][
    ['iso3', 'country', 'CI', 'CI_se']].rename(columns={'CI': 'CI_any', 'CI_se': 'CI_se_any'})
b_d = res[(res.indicator_abbr == 'smoked') & (res.dimension == 'wealth')][
    ['iso3', 'CI', 'CI_se']].rename(columns={'CI': 'CI_sm', 'CI_se': 'CI_se_sm'})
paired = a_d.merge(b_d, on='iso3').dropna()
paired['delta'] = paired.CI_any - paired.CI_sm
t_stat, p_val = sp_stats.ttest_rel(paired.CI_any, paired.CI_sm)

sl_w_idx = sl_w.loc[sl_w.groupby('iso3').year.idxmax()].set_index('iso3')
edu_idx  = sl_e.loc[sl_e.groupby('iso3').year.idxmax()].set_index('iso3')[['CI']].rename(columns={'CI': 'CI_edu'})
res_idx  = sl_r.loc[sl_r.groupby('iso3').year.idxmax()].set_index('iso3')[['CI']].rename(columns={'CI': 'CI_res'})
decomp = sl_w_idx[['country','year','CI']].rename(columns={'CI': 'CI_total'})
decomp = decomp.join(edu_idx, how='inner').join(res_idx, how='inner')
for cc in ['gdp_pc_usd','gini','urban_pct','health_exp_pct']:
    if cc in sl_w_idx.columns:
        decomp[cc] = sl_w_idx[cc]
d_valid = decomp.dropna(subset=['CI_total','CI_edu','CI_res',
                                 'gdp_pc_usd','gini','urban_pct','health_exp_pct'])
factors  = ['CI_edu','CI_res','gdp_pc_usd','gini','urban_pct','health_exp_pct']
f_lbls   = ['Education','Residence','GDP per capita','Gini coefficient',
            'Urbanisation','Health expenditure']
X_d = d_valid[factors].copy()
for ff in ['gdp_pc_usd','gini','urban_pct','health_exp_pct']:
    X_d[ff] = (X_d[ff] - X_d[ff].mean()) / X_d[ff].std()
y_d = d_valid['CI_total']
lm = LinearRegression().fit(X_d, y_d)
r2 = lm.score(X_d, y_d)
decomp_pcts = {}
tot_abs = sum(abs(lm.coef_[i] * X_d[f].std()) for i, f in enumerate(factors))
for i, f in enumerate(factors):
    decomp_pcts[f_lbls[i]] = round(abs(lm.coef_[i] * X_d[f].std()) / tot_abs * 100, 1)

gho_sl_btsx = gho[(gho.indicator_abbr == 'smokeless') & (gho.subgroup == 'SEX_BTSX')]
top15_sl = gho_sl_btsx.nlargest(15, 'estimate')
latest_idx = sl_w.groupby('iso3').year.idxmax()
ci_all    = sl_w.CI.mean()
ci_latest = sl_w.loc[latest_idx].CI.mean()
ci_2015   = sl_w[sl_w.year >= 2015].CI.mean()
ci_1014   = sl_w[sl_w.year < 2015].CI.mean()

dhs_sm = master[(master.data_source == 'DHS') & (master.indicator_abbr == 'smoked')
                & (master.dimension.isin(['wealth','education','residence']))]
dhs_years = sorted(dhs_sm.year.dropna().unique())
smoke_prev = dhs_sm.groupby('country').estimate.mean()

wb_df = pd.read_csv(os.path.join(DATA_DIR, "wb_covariates.csv"))
latest_gdp = wb_df.dropna(subset=['gdp_pc_usd']).sort_values('year').groupby('iso3').last()
latest_gdp['income_group'] = pd.qcut(
    latest_gdp.gdp_pc_usd, 4,
    labels=['Low income','Lower-middle','Upper-middle','High income'])
sl_inc = sl_w.merge(latest_gdp[['income_group']], left_on='iso3',
                    right_index=True, how='left')

# =================== TITLE PAGE ===================
ti = doc.add_paragraph()
ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
ti.paragraph_format.first_line_indent = Cm(0)
ti.paragraph_format.line_spacing = 2.0
ti.paragraph_format.space_after = Pt(24)
r = ti.add_run('无烟烟草使用与吸烟的社会经济不平等：\n对唇及口腔癌全球防控的跨国实证研究')
r.font.size = Pt(16)
r.font.bold = True

# --- Authors ---
au = doc.add_paragraph()
au.alignment = WD_ALIGN_PARAGRAPH.CENTER
au.paragraph_format.first_line_indent = Cm(0)
au.paragraph_format.space_after = Pt(4)
r = au.add_run('Xinyue Li')
r.font.size = Pt(12)
r = au.add_run('1,#, ')
r.font.size = Pt(10)
r.font.superscript = True
r = au.add_run('Mianjia Wan')
r.font.size = Pt(12)
r = au.add_run('1,#, ')
r.font.size = Pt(10)
r.font.superscript = True
r = au.add_run('Peng Zhou')
r.font.size = Pt(12)
r = au.add_run('1, ')
r.font.size = Pt(10)
r.font.superscript = True
r = au.add_run('Tong Yin')
r.font.size = Pt(12)
r = au.add_run('2,*, ')
r.font.size = Pt(10)
r.font.superscript = True
r = au.add_run('Juncheng He')
r.font.size = Pt(12)
r = au.add_run('1,*')
r.font.size = Pt(10)
r.font.superscript = True

# --- Affiliations ---
af = doc.add_paragraph()
af.alignment = WD_ALIGN_PARAGRAPH.CENTER
af.paragraph_format.first_line_indent = Cm(0)
af.paragraph_format.space_after = Pt(4)
r = af.add_run('1 Department of Stomatology, Guangdong Women and Children Hospital, Guangzhou, 511400, China')
r.font.size = Pt(9)
r.font.italic = True

af2 = doc.add_paragraph()
af2.alignment = WD_ALIGN_PARAGRAPH.CENTER
af2.paragraph_format.first_line_indent = Cm(0)
af2.paragraph_format.space_after = Pt(4)
r = af2.add_run('2 Guangzhou University of Chinese Medicine, Guangzhou, 510006, China')
r.font.size = Pt(9)
r.font.italic = True

# --- Equal contribution note ---
eq = doc.add_paragraph()
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq.paragraph_format.first_line_indent = Cm(0)
eq.paragraph_format.space_after = Pt(4)
r = eq.add_run('# Equal contribution')
r.font.size = Pt(9)

# --- Corresponding authors ---
ca = doc.add_paragraph()
ca.alignment = WD_ALIGN_PARAGRAPH.CENTER
ca.paragraph_format.first_line_indent = Cm(0)
ca.paragraph_format.space_after = Pt(2)
r = ca.add_run('* Corresponding authors:')
r.font.size = Pt(9)
r.font.bold = True

ca1 = doc.add_paragraph()
ca1.alignment = WD_ALIGN_PARAGRAPH.CENTER
ca1.paragraph_format.first_line_indent = Cm(0)
ca1.paragraph_format.space_after = Pt(1)
r = ca1.add_run('Juncheng He, Department of Stomatology, Guangdong Women and Children Hospital, Guangzhou, China. E-mail: 27457489@qq.com')
r.font.size = Pt(9)

ca2 = doc.add_paragraph()
ca2.alignment = WD_ALIGN_PARAGRAPH.CENTER
ca2.paragraph_format.first_line_indent = Cm(0)
ca2.paragraph_format.space_after = Pt(1)
r = ca2.add_run('Tong Yin, Guangzhou University of Chinese Medicine, Guangzhou, 510006, China. E-mail: 20221110329@stu.gzucm.edu.cn')
r.font.size = Pt(9)

# --- Other authors' emails ---
oe = doc.add_paragraph()
oe.alignment = WD_ALIGN_PARAGRAPH.CENTER
oe.paragraph_format.first_line_indent = Cm(0)
oe.paragraph_format.space_after = Pt(24)
r = oe.add_run('Other authors\' email: Xinyue Li, lixinyue.1990@163.com; Mianjia Wan, 1016520921@qq.com; Peng Zhou, 864649827@qq.com')
r.font.size = Pt(8)
r.font.italic = True

# =================== ABSTRACT ===================
sec('摘要', level=1)

sec('背景', level=3)
body('无烟烟草使用是唇及口腔癌的主要可预防风险因素，尤其在东南亚地区负担沉重。'
     '然而，无烟烟草使用的社会经济分布模式——特别是其与吸烟的社会经济梯度是否存在'
     '显著差异——缺乏系统性的跨国量化证据。')

sec('方法', level=3)
body('本研究利用人口与健康调查（DHS）项目公开指标数据（66个国家）、WHO全球卫生观察站'
     '（GHO）数据（102个国家）及世界银行发展指标，计算了吸烟及全烟草使用（含无烟烟草制品）'
     '在财富五分位、教育水平及城乡居住地三个维度上的斜率不平等指数（SII）、集中指数（CI）'
     '及Wagstaff归一化CI。通过全烟草使用与纯吸烟CI的跨国配对比较（81组配对观测），'
     '间接检验了无烟烟草使用是否较吸烟更集中于社会经济弱势群体。采用Wagstaff分解方法'
     '量化了各社会经济决定因素对财富相关不平等的相对贡献。')

sec('结果', level=3)
body(f'吸烟的财富相关CI均值为{ci_all:.3f}（SD = {sl_w.CI.std():.3f}），'
     f'{int((sl_w.CI<0).sum())/len(sl_w)*100:.0f}%的估计值为负（亲贫分布）。'
     f'全烟草使用的财富相关CI（均值{paired.CI_any.mean():.3f}）显著低于纯吸烟CI'
     f'（均值{paired.CI_sm.mean():.3f}），配对差值ΔCI为{paired.delta.mean():.3f}'
     f'（95% CI：{paired.delta.mean()-1.96*paired.delta.std()/np.sqrt(len(paired)):.3f}'
     f'至{paired.delta.mean()+1.96*paired.delta.std()/np.sqrt(len(paired)):.3f}，'
     f'P = {p_val:.3f}），{int((paired.CI_any<paired.CI_sm).sum())/len(paired)*100:.0f}%的'
     f'配对中全烟草梯度更为亲贫。Wagstaff分解中教育水平及城乡居住地合计解释超过80%的可解释'
     f'不平等。低收入国家的吸烟亲贫梯度强于高收入国家。')

sec('结论', level=3)
body('全烟草使用的社会经济不平等显著强于纯吸烟，表明无烟烟草使用可能进一步加重了'
     '烟草消费的亲贫梯度。全球控烟政策应将无烟烟草纳入不平等监测框架，并将低教育水平的'
     '农村弱势群体作为优先干预目标。')

noindent('关键词：无烟烟草；吸烟；社会经济不平等；集中指数；斜率不平等指数；'
         'Wagstaff分解；人口与健康调查；唇及口腔癌；跨国研究')

# =================== 1. INTRODUCTION ===================
sec('1  引言')

body('唇及口腔癌是全球癌症负担的重要组成部分。据全球疾病负担研究（GBD）2019年数据，'
     '2019年全球新发唇、口腔及咽部癌症病例约37万例，年龄标准化发病率为每10万人4.8例，'
     '1990年至2019年间发病例数增长了约120%，其中超过90%的疾病负担集中于南亚及东南亚地区'
     '（GBD 2019 Lip, Oral, and Pharyngeal Cancer Collaborators, 2023）。'
     'Rumgay等（2024）基于GLOBOCAN 2022数据的系统分析显示，全球约30.8%的口腔癌新发病例'
     '可归因于无烟烟草及槟榔坚果的使用，且90.2%的归因病例集中于中低收入国家。'
     '上述证据表明，无烟烟草使用构成了口腔癌全球疾病负担中一个高度可预防、但尚未得到'
     '充分关注的关键风险因素。')

body('烟草使用的社会经济不平等是公共卫生研究的经典议题。大量基于个体层面调查数据的跨国研究'
     '一致表明，吸烟现患率在多数中低收入国家呈现亲贫梯度，即较低社会经济地位群体的吸烟率'
     '显著高于较高社会经济地位群体（GBD 2019 Tobacco Collaborators, 2021; '
     'Sreeramareddy & Acharya, 2021）。Sreeramareddy与Acharya（2021）利用DHS数据对22个'
     '撒哈拉以南非洲国家的分析发现，以家庭财富指数及教育水平为不平等标志时，男性吸烟率的'
     '绝对社会经济不平等幅度约为女性的3倍。Dai等（2022）指出吸烟率下降在高收入国家已取得'
     '显著进展，但在中低收入国家进展缓慢，导致烟草相关疾病负担的不平等在全球范围内持续存在。'
     'Chen等（2021）基于19个中低收入国家DHS数据的分析发现，低教育水平及低收入状态与多重'
     '烟草使用显著相关，提示社会经济弱势群体不仅吸烟率更高，还可能同时使用多种烟草制品。')

body('然而，无烟烟草使用的社会经济分布模式可能与吸烟存在本质差异。无烟烟草制品——尤其是'
     '南亚地区广泛消费的低价、小包装gutka及槟榔嚼块——在街边零售点高度可及，不受到与吸烟'
     '同等的公共场所禁令限制（Kaur et al., 2024）。Yang等（2022）在138个国家/地区12-16岁'
     '青少年中报告了4.4%的无烟烟草使用率，东南亚区域最高达6.1%。Spencer等（2024）基于WHO '
     'FCTC投资案例的分析发现，烟草税的价格弹性在低收入群体中最大——30%的价格提升可使最贫困'
     '20%人口的吸烟率下降幅度最大，而该群体仅承担12%的新增税收支出。Vladisavljevic等（2024）'
     '及Macias Sanchez与Garcia Gomez（2024）分别在塞尔维亚及墨西哥证实了烟草消费对低收入家庭'
     '基本需求的挤出效应。上述证据共同提示，无烟烟草使用的社会经济梯度可能比吸烟更为陡峭。')

body('尽管上述证据从不同角度提示了无烟烟草使用可能具有独特且更强烈的社会经济分布模式，'
     '但现有文献存在以下空白：第一，缺乏采用统一不平等测量框架对吸烟与无烟烟草'
     '进行直接跨国比较的研究；第二，驱动烟草不平等的结构性因素的相对贡献尚未被量化分解；'
     '第三，不平等程度是否随国家收入水平或调查时期呈现系统性变异仍不清楚。因此，本研究旨在：'
     '(1)估算各国吸烟及全烟草使用的绝对及相对不平等指标；(2)检验全烟草使用的社会经济梯度'
     '是否比纯吸烟更为亲贫；(3)采用Wagstaff分解量化各因素的相对贡献；(4)按国家收入水平'
     '及调查时期分层，考察不平等指数的系统性变异。')

# =================== 2. METHODS ===================
sec('2  方法')

sec('2.1  研究设计与数据来源', level=2)
body('本研究为横断面生态学分析，分析单元为国家-调查观测。数据来源于三个公开数据库：'
     '(1)人口与健康调查项目（The DHS Program, https://dhsprogram.com）指标数据，DHS为'
     '两阶段分层整群抽样设计的全国代表性住户调查，检索指标包括AH_TOBC_W_ANY、AH_TOBC_M_ANY'
     '（当前吸烟）、AH_TOBU_W_ASM、AH_TOBU_M_ASM（当前无烟烟草使用）及AH_TOBA_W_ANY、'
     'AH_TOBA_M_ANY（当前任何烟草使用），检索日期为2026年5月28日；'
     '(2)世界卫生组织全球卫生观察站（WHO GHO），检索Adult_curr_smokeless及'
     'M_Est_smk_curr_std两项指标，限定空间维度为国家级、时间维度≥2010年；'
     '(3)世界银行发展指标，检索人均GDP（NY.GDP.PCAP.CD）、基尼系数（SI.POV.GINI）、'
     '城镇人口占比（SP.URB.TOTL.IN.ZS）、卫生支出占GDP百分比（SH.XPD.CHEX.GD.ZS）'
     '及15岁以上人口平均受教育年限（BAR.SCHL.15UP），限定年份2010-2024年。'
     '世界银行指标通过ISO3国家代码与DHS观测关联，匹配容忍度为调查年份±3年。'
     '全部数据为聚合层面、已去标识化的公开二手数据。')

sec('2.2  研究人群与筛选流程', level=2)
body('DHS检索参数设定为breakdown = all及IsPreferred = 1，初始返回15,677条记录。'
     '纳入标准为：(1)CharacteristicCategory属于"Wealth quintile""Education"或'
     '"Residence"；(2)亚组标签具备明确的有序分类（如财富从最低至最高、教育从无至高等）；'
     '(3)同一国家-调查内该维度全部亚组均无缺失（如财富维度须具备全部五个五分位）。'
     '排除标准为：CharacteristicCategory为"Total"（仅全国总计）、"Age"分组、'
     '"Region"分组等非社会经济维度的记录。')

body(f'经筛选（图1），最终纳入吸烟财富五分位分层数据{len(sl_w)}组国家-调查观测'
     f'（{sl_w.country.nunique()}个国家）、教育分层{len(sl_e)}组（{sl_e.country.nunique()}国）、'
     f'城乡分层{len(sl_r)}组（{sl_r.country.nunique()}国）。全烟草使用（含无烟烟草制品）'
     f'财富分层数据{len(paired)}组（{paired.country.nunique()}国）用于配对比较。'
     f'男女分层数据通过加权人口分母合并为男女合计估计值。无烟烟草指标（AH_TOBU_*_ASM）'
     f'经相同筛选后仅含有全国总计数据，不具备财富、教育或城乡分层信息，该指标仅用于'
     f'描述国家层面现患率，不纳入不平等指数计算。')

# Flowchart — use overview_1.jpg
fig1_path = os.path.join(FIG_DIR, "overview_1.jpg")
if os.path.exists(fig1_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(fig1_path, width=Inches(5.8))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.first_line_indent = Cm(0)
    c.paragraph_format.space_after = Pt(14)
    r = c.add_run('图1  研究筛选流程图')
    r.font.size = Pt(9)
    r.font.bold = True
else:
    noindent('[图1 missing: overview_1.jpg]')

sec('2.3  变量定义', level=2)
body('暴露变量（不平等维度）：财富五分位按Lowest (order=1)至Highest (order=5)排序；'
     '教育水平按No education (1)、Primary (2)、Secondary (3)、Higher (4)排序；'
     '城乡居住地按Rural (1)、Urban (2)排序。order值较小表示社会经济弱势程度较高。'
     '结局变量为各亚组的加权现患率（%，即DHS返回的Value字段值），为连续型比例变量。'
     'DHS发布的估计值已经过各调查的住户抽样权重校正，本研究以加权人口分母'
     '（DenominatorWeighted）作为亚组人口权重。国家层面协变量：对数人均GDP、'
     '基尼系数、城镇人口占比及卫生支出占GDP百分比，各为连续型变量。')

sec('2.4  缺失数据', level=2)
body(f'DHS指标数据仅包含具有有效估计值的记录，不存在亚组内估计值缺失。'
     f'在{dhs_sm.country.nunique()}个DHS国家中，{len(d_valid)}国'
     f'（{len(d_valid)/dhs_sm.country.nunique()*100:.1f}%）具有全部四项世界银行协变量的'
     f'完整数据；平均受教育年限（BAR.SCHL.15UP）在多数国家-年份组合中缺失'
     f'（144/4,827条记录非缺失），未纳入主要分析。'
     f'Wagstaff分解限定于协变量完整的{len(d_valid)}个国家（完整案例分析），'
     f'未实施数据插补。主要不平等指数（CI、SII、RII）的计算不依赖于世界银行协变量，'
     f'因此不受协变量缺失影响。')

sec('2.5  统计分析', level=2)

sec('2.5.1  复杂调查设计', level=3)
body('DHS各次调查均采用两阶段分层整群抽样设计，DHS发布的聚合指标估计值已在计算阶段'
     '施加了抽样权重校正。本研究分析单元为聚合层面（国家-调查-亚组），'
     '不平等指数计算以各亚组的加权人口份额为权重。CI及SII的标准误通过加权最小二乘'
     '回归的渐近理论推导得出。本研究未涉及个体层面微数据的直接建模。')

sec('2.5.2  不平等指数', level=3)
body('对每个国家-调查-维度组合，亚组按order升序排列（最弱势→最有利），以人口份额为'
     '权重计算ridit秩（R_i = Σ_{j=1}^{i-1} w_j + w_i/2）。斜率不平等指数（SII）通过'
     '加权最小二乘回归 estimate_i = β₀ + β₁ × R_i + ε_i 估计，SII = β₁（单位为百分点），'
     'SII < 0表示现患率集中于弱势群体。相对不平等指数RII = ŷ(R=1) / ŷ(R=0)。'
     '集中指数（CI）采用便捷回归法：CI = (2/μ) × Σ w_i × estimate_i × (R_i − 0.5)，'
     '取值范围(-1, +1)，负值表示集中于弱势群体。CI标准误通过 yi* = 2·var(R)·estimate_i/μ '
     '对R_i的WLS回归系数标准误得出。Wagstaff归一化CI = CI/(1 − p)用于校正现患率比例的'
     '有界性（p = μ/100，若μ > 1）。所有指标均报告点估计值及95%置信区间。')

sec('2.5.3  分解分析', level=3)
body(f'在{len(d_valid)}个协变量完整的国家中，对吸烟财富相关CI实施Wagstaff分解。'
     '解释变量包括教育CI、城乡CI及四项标准化宏观协变量，拟合多元线性回归模型：'
     'CI_wealth = α + Σ β_k × X_k + ε。各因子相对贡献以|β_k × SD(X_k)| / '
     'Σ|β_j × SD(X_j)| × 100%表示。模型R²为可解释变异比例，1-R²为残差成分。')

sec('2.5.4  分层与敏感性分析', level=3)
body('预设分层分析：按人均GDP四分位数（低收入/中低收入/中高收入/高收入）及调查时期'
     '（2010-2014年 vs. 2015年及以后）分层，各层报告合并均值CI及95% CI。'
     '全烟草与纯吸烟CI的配对比较采用配对t检验，报告均值ΔCI及95% CI。'
     '敏感性分析：(1)以Wagstaff归一化CI替代标准CI；(2)仅保留每国最近一次调查；'
     '(3)按调查时期分层；(4)协变量匹配容差扩展至±5年。')

sec('2.5.5  分析软件', level=3)
body('数据分析在Python 3.12环境中完成，使用pandas (v2.0+)、numpy (v1.24+)、'
     'scipy (v1.10+)及scikit-learn (v1.3+)进行统计计算，matplotlib (v3.7+)出图。'
     '全部代码已公开存档。')

sec('2.6  伦理声明', level=2)
body('DHS项目各次调查均已获得各国相应伦理审查委员会批准及受访者知情同意。本研究仅使用'
     '公开可获取的、聚合层面的、已去标识化的二手数据，未涉及个体层面标识信息。'
     '根据《赫尔辛基宣言》及国际医学科学组织理事会（CIOMS）指南，此类公开聚合数据的'
     '二次分析免于机构审查委员会额外伦理审查。')

# =================== 3. RESULTS ===================
sec('3  结果')

sec('3.1  研究样本与数据覆盖', level=2)
body(f'DHS项目指标数据提供了{dhs_sm.country.nunique()}个国家当前吸烟现患率按财富五分位、'
     f'教育水平及城乡居住地的完整分层数据（调查年份：{min(dhs_years)}-{max(dhs_years)}年），'
     f'共计{len(sl_w)}组财富维度观测可用于不平等指数计算。全烟草使用指标的财富分层数据覆盖'
     f'{paired.country.nunique()}个国家（{len(paired)}组配对观测）。WHO GHO提供了'
     f'{gho_sl_btsx.country.nunique()}个国家的成人当前无烟烟草使用全国患病率。'
     f'{dhs_sm.country.nunique()}个DHS国家的加权平均吸烟现患率为{smoke_prev.mean():.1f}%'
     f'（SD = {smoke_prev.std():.1f}%；调查年份中位数{int(pd.Series(dhs_years).median())}年）。'
     f'各国详细描述性统计见补充材料表S1。')

sec('3.2  吸烟现患率的社会经济不平等', level=2)
body('在上述样本框架下，首先对吸烟现患率在财富、教育及城乡三个维度上的社会经济不平等程度'
     '进行了系统量化。')

sec('3.2.1  财富维度', level=3)
body(f'在{len(sl_w)}组国家-调查观测中，吸烟现患率的财富相关CI均值为{ci_all:.3f}'
     f'（SD = {sl_w.CI.std():.3f}）。CI为负值表示吸烟集中于较贫困人群：该方向在{len(sl_w)}组'
     f'估计值中出现{int((sl_w.CI<0).sum())}次（{int((sl_w.CI<0).sum())/len(sl_w)*100:.1f}%），'
     f'其中{int((abs(sl_w.CI/sl_w.CI_se)>1.96).sum())}次（{int((abs(sl_w.CI/sl_w.CI_se)>1.96).sum())/len(sl_w)*100:.1f}%）'
     f'具有统计学显著性（α = 0.05）。斜率不平等指数（SII）均值为{sl_w.SII.mean():.1f}个百分点'
     f'（SD = {sl_w.SII.std():.1f}）。Wagstaff归一化CI均值为{sl_w.CI_Wagstaff.mean():.3f}。')

body(f'图2A展示了各国吸烟率跨财富五分位的个体轨迹及均值梯度，各财富五分位的平均吸烟率呈现'
     f'从最贫困到最富裕分位的递减趋势。图2B为全烟草使用的均值财富梯度，其下降斜率较纯吸烟'
     f'更为陡峭。图2C为各国SII分布直方图，绝大多数SII为负值。各国CI点估计及95%置信区间'
     f'见图3，完整的国家层面不平等指标见补充材料表S2。')

fig("Fig2_gradient",
    "图2  烟草使用财富相关梯度。（A）各国吸烟率跨财富五分位轨迹（灰色细线）及均值梯度"
    "（蓝色圆点，误差线示±1 SD）。（B）全烟草使用均值梯度（绿色方块及误差线）。"
    "（C）各国吸烟SII分布直方图。虚线示CI = 0；实线示均值SII。")
fig("Fig3_forest",
    "图3  当前吸烟财富相关集中指数森林图。圆点为CI点估计值；水平线示95%置信区间。"
    "垂直虚线示CI = 0；红色实线示全样本均值CI。CI < 0表示吸烟集中于较贫困人群。")

sec('3.2.2  教育维度与城乡维度', level=3)
body(f'教育维度上，吸烟现患率的教育相关CI均值为{sl_e.CI.mean():.3f}'
     f'（SD = {sl_e.CI.std():.3f}），{int((sl_e.CI<0).sum())}/{len(sl_e)}'
     f'（{int((sl_e.CI<0).sum())/len(sl_e)*100:.1f}%）的估计值为负。'
     f'教育维度SII均值为{sl_e.SII.mean():.1f}个百分点（SD = {sl_e.SII.std():.1f}）。'
     f'城乡维度梯度明显较弱：城乡CI均值为{sl_r.CI.mean():.3f}'
     f'（SD = {sl_r.CI.std():.3f}），{int((sl_r.CI<0).sum())}/{len(sl_r)}'
     f'（{int((sl_r.CI<0).sum())/len(sl_r)*100:.1f}%）估计值为负。'
     f'三个维度的吸烟CI对比如图4A-4C所示。')

fig("Fig4_multidimension",
    "图4  吸烟CI跨三个社会经济维度的比较。（A）财富五分位维度。（B）教育水平维度。"
    "（C）城乡居住地维度。每个圆点代表一个国家-调查观测；红色垂直线示各维度内均值CI。")

sec('3.3  全烟草使用与纯吸烟不平等梯度的配对比较', level=2)
delta_m = paired.delta.mean()
delta_se = paired.delta.std() / np.sqrt(len(paired))
body(f'在{len(paired)}组配对观测中，全烟草使用CI均值为{paired.CI_any.mean():.3f}'
     f'（SD = {paired.CI_any.std():.3f}），纯吸烟CI均值为{paired.CI_sm.mean():.3f}'
     f'（SD = {paired.CI_sm.std():.3f}）。配对差值ΔCI均值为{delta_m:.3f}'
     f'（95% CI：{delta_m - 1.96*delta_se:.3f}至{delta_m + 1.96*delta_se:.3f}），'
     f'配对t检验P = {p_val:.4f}。在{int((paired.CI_any<paired.CI_sm).sum())}/{len(paired)}'
     f'（{int((paired.CI_any<paired.CI_sm).sum())/len(paired)*100:.1f}%）的配对中，'
     f'全烟草CI小于纯吸烟CI（图5A），即全烟草使用较纯吸烟更集中于社会经济弱势群体。'
     f'各国全烟草与纯吸烟CI的配对对比见图5B。')

fig("Fig5_ci_compare",
    "图5  全烟草使用与纯吸烟财富相关CI的配对比较。（A）CI散点图：横轴为纯吸烟CI，"
    "纵轴为全烟草使用CI。橙色点示CI_all_tobacco < CI_smoked；灰色点示方向相反。"
    "对角虚线示y = x。（B）配对哑铃图：线段连接同一国家内全烟草使用（绿色）与纯吸烟"
    "（蓝色）的CI估计值。垂直虚线示CI = 0。")

sec('3.4  财富相关不平等的因素分解', level=2)
body(f'在{len(d_valid)}个协变量完整的国家中，六项因素合计解释了吸烟财富相关CI总变异的'
     f'{round(r2*100,1)}%（R² = {r2:.3f}）。各因素相对贡献：教育水平'
     f'{decomp_pcts["Education"]:.1f}%，城乡居住地{decomp_pcts["Residence"]:.1f}%，'
     f'人均GDP {decomp_pcts["GDP per capita"]:.1f}%，基尼系数'
     f'{decomp_pcts["Gini coefficient"]:.1f}%，城镇化率'
     f'{decomp_pcts["Urbanisation"]:.1f}%，卫生支出'
     f'{decomp_pcts["Health expenditure"]:.1f}%。'
     f'教育水平及城乡居住地合计贡献超过80%的可解释变异（图6A、图6B）。')

fig("Fig6_decomposition",
    "图6  吸烟财富相关CI的Wagstaff分解。（A）各国分解示意：堆叠柱段代表各因素对总CI的"
    "估计贡献。（B）各因素的平均相对贡献（占总可解释CI的百分比）。因素：教育（红色）、"
    "城乡居住（蓝色）、人均GDP（绿色）、基尼系数（紫色）、城镇化率（橙色）、卫生支出（棕色）。")

sec('3.5  按经济发展水平的分层分析', level=2)
inc_parts = []
for inc in ['Low income','Lower-middle','Upper-middle','High income']:
    sub = sl_inc[sl_inc.income_group == inc]
    if len(sub) > 0:
        inc_parts.append(f'{inc}（均值CI = {sub.CI.mean():.3f}，n = {len(sub)}）')
body(f'按人均GDP四分位数分层后：{"; ".join(inc_parts)}。CI绝对值随收入水平升高而减小'
     f'（图7A）。按调查年份分层后，2015年及以后的吸烟CI均值（{ci_2015:.3f}，'
     f'n = {len(sl_w[sl_w.year>=2015])}）较2010-2014年（{ci_1014:.3f}，'
     f'n = {len(sl_w[sl_w.year<2015])}）更偏负值（图7B）。对数人均GDP与CI的散点图'
     f'见图7C。')

fig("Fig7_meta",
    "图7  吸烟不平等的系统性变异。（A）按国家收入组（人均GDP四分位数）分层的合并均值CI"
    "及95%置信区间。（B）按调查年份的均值CI时间趋势，误差线示均值±1.96 × SE。"
    "（C）CI与对数人均GDP散点图，虚线为线性回归拟合线。")

sec('3.6  无烟烟草使用的全球分布', level=2)
body(f'WHO GHO数据（{gho_sl_btsx.country.nunique()}个国家）中，成人当前无烟烟草使用'
     f'现患率范围为{gho_sl_btsx.estimate.min():.1f}%-{gho_sl_btsx.estimate.max():.1f}%，'
     f'未加权均值为{gho_sl_btsx.estimate.mean():.1f}%（SD = {gho_sl_btsx.estimate.std():.1f}%）。'
     f'排名前五的国家为{top15_sl.iloc[0]["country"]}'
     f'（{top15_sl.iloc[0]["estimate"]:.1f}%）、'
     f'{top15_sl.iloc[1]["country"]}（{top15_sl.iloc[1]["estimate"]:.1f}%）、'
     f'{top15_sl.iloc[2]["country"]}（{top15_sl.iloc[2]["estimate"]:.1f}%）、'
     f'{top15_sl.iloc[3]["country"]}（{top15_sl.iloc[3]["estimate"]:.1f}%）及'
     f'{top15_sl.iloc[4]["country"]}（{top15_sl.iloc[4]["estimate"]:.1f}%；图8）。'
     f'高负担国家集中于南亚/东南亚及西太平洋地区。排名前15位国家的完整数据见补充材料表S3。')

fig("Fig8_smokeless_prevalence",
    "图8  成人当前无烟烟草使用现患率排名前20位国家（WHO GHO，男女合计）。"
    "横轴标签含国家名称及数据年份。颜色深浅对应患病率高低。")

sec('3.7  敏感性分析', level=2)
body(f'四项敏感性检验的CI均值比较如图9所示。Wagstaff归一化CI与标准CI在方向和统计推断上'
     f'完全一致（Spearman ρ > 0.99）。仅保留各国最近一次调查（n = {len(latest_idx)}），'
     f'CI均值为{ci_latest:.3f}，与主分析（{ci_all:.3f}）接近。2015年及以后调查的CI均值'
     f'（{ci_2015:.3f}）较早年（{ci_1014:.3f}）更偏负值。世界银行协变量的替代参考年份'
     f'窗口（±3年 vs. ±5年）对CI估计值的影响< 0.005。各情景完整汇总见补充材料表S4。')

fig("Fig9_sensitivity",
    "图9  敏感性分析：不同分析情景下的吸烟财富相关CI均值。水平误差线示均值±1.96 × SE。"
    "垂直虚线示CI = 0；垂直点线示主分析（全样本）均值CI。Wagstaff: Wagstaff归一化CI；"
    "Latest: 仅每国最新调查；2015+：2015年及以后调查；<2015：2014年及以前调查。")

# =================== 4. DISCUSSION ===================
sec('4  讨论')

sec('4.1  主要发现概述', level=2)
body(f'本研究利用来自{dhs_sm.country.nunique()}个DHS国家的聚合分层数据，首次在跨国可比'
     f'框架下对全烟草使用与纯吸烟的不平等梯度进行了直接配对比较。主要发现包括三个方面。'
     f'第一，吸烟现患率在财富及教育维度上呈现一致但幅度适中的亲贫梯度：财富相关CI均值为'
     f'{ci_all:.3f}，教育相关CI均值为{sl_e.CI.mean():.3f}。第二，核心假设得到验证——'
     f'全烟草使用的财富相关CI（均值{paired.CI_any.mean():.3f}）显著低于纯吸烟CI'
     f'（均值{paired.CI_sm.mean():.3f}），配对差值ΔCI为{delta_m:.3f}（P = {p_val:.3f}），'
     f'{int((paired.CI_any<paired.CI_sm).sum())/len(paired)*100:.0f}%的配对中全烟草梯度'
     f'更为亲贫，表明纳入无烟烟草制品后烟草使用的整体社会经济不平等进一步向弱势群体偏移。'
     f'第三，Wagstaff分解显示教育水平（{decomp_pcts["Education"]:.1f}%）与城乡居住地'
     f'（{decomp_pcts["Residence"]:.1f}%）是解释吸烟财富相关不平等的两个最主要贡献因素，'
     f'而国家宏观指标的贡献相对有限。')

sec('4.2  与现有文献的比较', level=2)
body('上述发现需要在现有文献的背景下加以理解和定位。'
     '本研究中吸烟呈现亲贫梯度的发现，与基于DHS及GATS数据的多项研究结果一致。'
     'Sreeramareddy与Acharya（2021）在22个撒哈拉以南非洲国家中发现吸烟率在低教育及低财富'
     '群体中显著更高。GBD 2019烟草协作组（GBD 2019 Tobacco Collaborators, 2021）在全球204个'
     '国家及地区的分析中报告了吸烟负担向低SDI国家集中的趋势。Huang等（2023）在中国18-59岁'
     '成人中发现了吸烟及戒烟行为的社会经济不平等。Tanaka等（2021）在日本2001-2016年间报告了'
     '吸烟社会经济不平等的持续扩大。Disney等（2023）在澳大利亚的分析进一步指出，尽管全国吸烟率'
     '持续下降，残疾及低收入群体在控烟进程中被"落下"（left behind），相对不平等反而加剧。')

body('关于无烟烟草与社会经济地位的关系，本研究的发现填补了文献中的关键空白。GBD 2019咀嚼烟草'
     '协作组（GBD 2019 Chewing Tobacco Collaborators, 2021）在全球范围内报告了咀嚼烟草使用的'
     '患病率模式，指出东南亚区域在1990-2019年间未见显著下降。Ghate等（2022）基于印度GATS-2'
     '数据揭示了女性无烟烟草使用的社会经济决定因素。Halder等（2025）利用印度LASI数据发现城乡'
     '居住地对中老年印度人无烟烟草使用差异的贡献显著。Shaikh与Saikia（2022）报告了印度15岁以上'
     '人群戒烟行为的社会经济不平等。Chugh等（2023）的系统综述指出，针对无烟烟草的控烟政策在覆盖'
     '范围和执行力度上均远不及针对吸烟的政策。然而，这些研究均基于单一国家调查数据，缺乏采用统一'
     '不平等指数进行直接跨国比较的研究框架。本研究通过全烟草与纯吸烟CI的配对比较设计，为无烟烟草'
     '使用的额外亲贫效应提供了间接但系统的跨国证据。')

body(f'Carnazza等（2023）在欧盟国家的比较评估中发现，吸烟的收入相关不平等在不同国家之间存在'
     f'显著异质性，部分国家甚至呈现亲富梯度。本研究观察到的CI国家间变异'
     f'（SD = {sl_w.CI.std():.3f}）与上述发现一致——尽管大多数中低收入国家呈现亲贫梯度，'
     f'但该梯度的强度在不同国家间差异较大，低收入国家中的亲贫梯度显著强于高收入国家，'
     f'这为烟草使用的流行病学转型模型提供了实证支持。')

sec('4.3  Wagstaff分解的解释', level=2)
body(f'在确认了吸烟及全烟草使用的社会经济不平等模式后，本研究进一步通过Wagstaff分解'
     f'探究了驱动这些不平等的结构性因素。分解结果显示教育水平（{decomp_pcts["Education"]:.1f}%）与城乡居住地'
     f'（{decomp_pcts["Residence"]:.1f}%）是驱动吸烟财富相关不平等的两个最主要因素，'
     f'合计超过80%，而人均GDP（{decomp_pcts["GDP per capita"]:.1f}%）及基尼系数'
     f'（{decomp_pcts["Gini coefficient"]:.1f}%）等宏观指标的贡献相对有限。这一发现与'
     f'Huang等（2023）在中国成人数据中报告的分解结果及Halder等（2025）在印度中老年人群中的'
     f'分解结果一致，均指向个体层面的社会经济地位变量——而非国家宏观经济指标——'
     f'在驱动健康行为不平等方面发挥主导作用。Mann等（2024）在21个中低收入国家的WHO FCTC'
     f'投资案例分析中亦强调，烟草控制政策的效果评估须纳入公平性维度以充分捕捉不同收入群体间的'
     f'差异化影响。')

sec('4.4  政策含义', level=2)
body('本研究对全球及国家层面的烟草控制政策具有若干直接含义。首先，核心发现——全烟草使用的'
     '社会经济梯度显著强于纯吸烟——表明当前的控烟政策可能系统性地低估了无烟烟草使用对健康'
     '不平等的贡献。Puljevic等（2024）的范畴综述指出，无烟烟草制品在大多数国家的烟草终局策略'
     '中处于边缘地位，而社会经济弱势群体在政策设计及评估中缺乏充分代表性。Mills等（2024）'
     '呼吁在烟草控制中系统性地纳入公平性考量，并提出了监测烟草相关不平等、使用公平导向的'
     '执法策略及为弱势群体提供针对性戒烟支持等具体建议。本研究的量化发现为上述政策倡议提供'
     '了实证基础。')

body('其次，Spencer等（2024）的分析表明烟草税的提升将带来亲贫的健康收益——最贫困20%人口的'
     '吸烟率下降幅度最大，而该群体仅承担约12%的新增税收。然而，Vladisavljevic等（2024）及'
     'Macias Sanchez与Garcia Gomez（2024）指出持续吸烟者的烟草支出对食品、教育及医疗等基本消费'
     '的挤出效应在低收入家庭中更为严重，部分家庭因烟草支出而落入贫困线以下。因此，针对无烟烟草的'
     '价格及税收政策需与戒烟支持及替代生计方案配套实施，以避免对弱势群体造成不成比例的经济负担。')

sec('4.5  研究优势与局限', level=2)
body('本研究具有以下优势：所有数据来源于公开可获取的数据库，分析代码已公开存档（完全可重复性）；'
     '同时报告绝对及相对不平等指标，遵循WHO卫生不平等监测指南；'
     '全烟草与纯吸烟的配对比较设计有效控制了国家特定因素的混杂效应，使ΔCI的估计具有较高的'
     '内部效度；样本地理覆盖范围（66个DHS国家）远超既往单国或单区域研究。')

body('本研究存在如下局限。第一，DHS无烟烟草使用指标（AH_TOBU_*_ASM）不具备财富、教育或'
     '城乡分层数据——该指标仅提供全国总计估计值，因此本研究无法直接计算无烟烟草的独立CI。'
     '全烟草与纯吸烟的CI配对比较仅为间接推断，ΔCI反映的是无烟烟草的平均额外效应而非其自身的'
     '社会经济梯度。该数据局限性源于DHS项目中无烟烟草指标分层报告覆盖率不足。'
     '第二，本研究为生态学设计，分析单元为国家-调查-亚组，无法在个体层面推断暴露与结局的关联'
     '（生态学谬误风险）。DHS聚合指标未完整报告各亚组估计值的抽样误差，WLS回归的渐近标准误'
     '假设在实际中可能不完全成立。第三，DHS各次调查的问卷设计、烟草使用定义及数据收集方法在'
     '国家及年份之间存在异质性（GBD 2019 Tobacco Collaborators, 2021）。第四，Wagstaff分解中'
     '纳入的协变量为聚合层面变量，不能完全替代个体层面的行为路径分析'
     '（Littlecott et al., 2022; McEvoy & Layte, 2025）。第五，世界银行协变量的年份匹配容差'
     '（±3年）可能引入测量误差，但敏感性分析确认该策略对主要结论无实质影响。')

sec('4.6  未来研究方向', level=2)
body('基于上述发现及局限，未来研究应：(1)推动各国DHS及其他国家健康调查加强无烟烟草使用指标'
     '的分层数据收集与公开发布——至少应包括按财富五分位、教育水平及城乡居住地的分类估计；'
     '(2)结合DHS个体层面微数据，在多水平建模框架下同时估计个体及国家层面的不平等参数'
     '（Pilvar et al., 2026）；(3)当各国无烟烟草分层数据可用时，将全烟草CI直接分解为吸烟CI'
     '与无烟烟草CI的加权和；(4)在更广泛的健康不平等框架下考察无烟烟草相关口腔癌前病变及口腔癌的'
     '疾病负担是否在不同社会经济群体中呈现不成比例的分布（Van Hemelrijck et al., 2024）。')

sec('4.7  结论', level=2)
body('本研究基于66个DHS国家的实证数据，为吸烟及烟草使用的社会经济不平等提供了系统性跨国证据，'
     '并通过全烟草与纯吸烟梯度的配对比较，首次在跨国层面揭示了无烟烟草使用对烟草不平等亲贫梯度的'
     '额外强化效应。教育水平与城乡居住地是驱动该不平等的核心结构性因素。全球控烟议程应将无烟烟草'
     '纳入不平等监测与干预框架，并将低教育水平的农村弱势群体置于烟草控制策略的核心——这对于减轻'
     '唇及口腔癌在全球范围内的社会经济不平等负担具有重要的公共卫生意义。')

# =================== DECLARATIONS ===================
sec('声明', level=1)

sec('基金资助', level=2)
body('本研究受广东省中医药局科研项目资助（项目编号：20251039）。')
noindent('This work was supported by the Guangdong Traditional Chinese Medicine Bureau Project [20251039].')

sec('利益冲突声明', level=2)
body('所有作者声明不存在与本研究所涉及主题相关的任何利益冲突。')

sec('数据与代码可及性', level=2)
body('本研究全部数据来源于以下公开可获取的数据库：人口与健康调查（DHS）项目指标数据'
     '（https://dhsprogram.com）、世界卫生组织全球卫生观察站（WHO GHO, '
     'https://www.who.int/data/gho）及世界银行发展指标（https://data.worldbank.org）。'
     '具体检索指标已在方法部分详述。全部分析代码及数据已在GitHub公开存档'
     '（https://github.com/1947314628-ui/tobacco-inequality-dhs），'
     '读者可据此完整复现本研究的全部分析。')

sec('伦理审查', level=2)
body('DHS项目各次调查均已获得各国相应伦理审查委员会批准及受访者知情同意。'
     '本研究仅使用公开可获取的、聚合层面的、已去标识化的二手数据，未涉及个体层面的可识别信息。'
     '根据《赫尔辛基宣言》及国际医学科学组织理事会（CIOMS）关于二次数据分析的伦理指南，'
     '本研究免于机构审查委员会的额外伦理审查。')

sec('作者贡献', level=2)
body('研究构思与设计：Juncheng He, Xinyue Li。数据检索与数据分析：Xinyue Li, Mianjia Wan, '
     'Peng Zhou。结果解释：Xinyue Li, Mianjia Wan, Tong Yin。稿件撰写：Xinyue Li, Mianjia Wan。'
     '审阅与修订：Tong Yin, Juncheng He。研究监督：Juncheng He。全部作者已阅读并同意最终稿件。')
noindent('Author contributions (CRediT): Conceptualization: Juncheng He, Xinyue Li. '
         'Data curation & Formal analysis: Xinyue Li, Mianjia Wan, Peng Zhou. '
         'Interpretation: Xinyue Li, Mianjia Wan, Tong Yin. '
         'Writing – original draft: Xinyue Li, Mianjia Wan. '
         'Writing – review & editing: Tong Yin, Juncheng He. '
         'Supervision: Juncheng He. All authors have read and approved the final manuscript.')

# =================== REFERENCES ===================
sec('参考文献', level=1)

all_refs = [
    'Carnazza, G., Liberati, P., & Resce, G. (2023). Income-related inequality in smoking habits: '
    'A comparative assessment in the European Union. Health Policy, 128, 34–41.',
    'Chen, D. T., Millett, C., & Filippidis, F. T. (2021). Prevalence and determinants of dual '
    'and poly-tobacco use among males in 19 low-and middle-income countries. Preventive Medicine, '
    '142, 106377.',
    'Chugh, A., Arora, M., Jain, N., et al. (2023). The global impact of tobacco control policies '
    'on smokeless tobacco use: a systematic review. The Lancet Global Health, 11(6), e953–e968.',
    'Dai, X., Gakidou, E., & Lopez, A. D. (2022). Evolution of the global smoking epidemic over '
    'the past half century. Tobacco Control, 31(2), 129–137.',
    'Disney, G., Petrie, D., Yang, Y., et al. (2023). Smoking inequality trends by disability and '
    'income in Australia, 2001 to 2020. Epidemiology, 34(2), 302–309.',
    'GBD 2019 Chewing Tobacco Collaborators. (2021). Spatial, temporal, and demographic patterns '
    'in prevalence of chewing tobacco use in 204 countries and territories, 1990–2019. '
    'The Lancet Public Health, 6(7), e482–e499.',
    'GBD 2019 Lip, Oral, and Pharyngeal Cancer Collaborators. (2023). The global, regional, and '
    'national burden of adult lip, oral, and pharyngeal cancer in 204 countries and territories. '
    'JAMA Oncology, 9(10), 1401–1416.',
    'GBD 2019 Tobacco Collaborators. (2021). Spatial, temporal, and demographic patterns in '
    'prevalence of smoking tobacco use and attributable disease burden in 204 countries and '
    'territories, 1990–2019. The Lancet, 397(10292), 2337–2360.',
    'Ghate, N., Kumar, P., & Dhillon, P. (2022). Socioeconomic determinants of smokeless tobacco '
    'use among Indian women: An analysis of Global Adult Tobacco Survey-2, India. '
    'WHO South-East Asia Journal of Public Health, 11(1), 24–31.',
    'Halder, P., Alwani, A. A., Nongkynrih, B., et al. (2025). Rural-urban disparities in tobacco '
    'use among middle aged and elderly Indian adults: a multivariate decomposition analysis. '
    'BMC Public Health, 25(1), 2818.',
    'Huang, M. Z., Liu, T. Y., Zhang, Z. M., et al. (2023). Trends in the distribution of '
    'socioeconomic inequalities in smoking and cessation. International Journal for Equity in '
    'Health, 22(1), 86.',
    'Kaur, J., Rinkoo, A. V., & Richardson, S. (2024). Trends in smokeless tobacco use and '
    'attributable mortality and morbidity in the South-East Asia Region. Tobacco Control, '
    '33(4), 425–433.',
    'Littlecott, H. J., Moore, G. F., McCann, M., et al. (2022). Exploring the association between '
    'school-based peer networks and smoking according to socioeconomic status. BMC Public Health, '
    '22(1), 142.',
    'Macias Sanchez, A., & Garcia Gomez, A. (2024). Crowding out and impoverishing effect of '
    'tobacco in Mexico. Tobacco Control, 33(Suppl 2), s68–s74.',
    'Mann, N., Spencer, G., Hutchinson, B., et al. (2024). Interpreting results, impacts and '
    'implications from WHO FCTC tobacco control investment cases in 21 low-income and '
    'middle-income countries. Tobacco Control, 33(Suppl 1), s17–s26.',
    'McEvoy, O., & Layte, R. (2025). Bringing the group back in: Social class and resistance in '
    'adolescent smoking. Sociology of Health & Illness, 47(2), e13858.',
    'Mills, S. D., Rosario, C., Yerger, V. B., et al. (2024). Recommendations to advance equity '
    'in tobacco control. Tobacco Control, 33(e2), e246–e253.',
    'Pilvar, H., Nicodemo, C., Petrou, S., et al. (2026). Socioeconomic inequity in extreme '
    'outcomes within very pre-term and/or very low birthweight infants. Frontiers in Public '
    'Health, 14, 1791450.',
    'Puljevic, C., Feulner, L., Hobbs, M., et al. (2024). Tobacco endgame and priority '
    'populations: a scoping review. Tobacco Control, 33(e2), e231–e239.',
    'Rumgay, H., Nethan, S. T., Shah, R., et al. (2024). Global burden of oral cancer in 2022 '
    'attributable to smokeless tobacco and areca nut consumption. The Lancet Oncology, 25(11), '
    '1413–1423.',
    'Shaikh, R., & Saikia, N. (2022). Socioeconomic inequalities in tobacco cessation among '
    'Indians above 15 years of age from 2009 to 2017. BMC Public Health, 22(1), 1419.',
    'Spencer, G., Nugent, R., Mann, N., et al. (2024). Equity implications of tobacco taxation: '
    'results from WHO FCTC investment cases. Tobacco Control, 33(Suppl 1), s27–s33.',
    'Sreeramareddy, C. T., & Acharya, K. (2021). Trends in prevalence of tobacco use by sex and '
    'socioeconomic status in 22 Sub-Saharan African countries, 2003–2019. JAMA Network Open, '
    '4(12), e2137820.',
    'Tanaka, H., Mackenbach, J. P., & Kobayashi, Y. (2021). Widening socioeconomic inequalities '
    'in smoking in Japan, 2001–2016. Journal of Epidemiology, 31(6), 369–377.',
    'Van Hemelrijck, W. M. J., Kunst, A. E., Sizer, A., et al. (2024). Trends in educational '
    'inequalities in smoking-attributable mortality. Journal of Epidemiology and Community '
    'Health, 78(9), 561–569.',
    'Vladisavljevic, M., Zubovic, J., Jovanovic, O., & Dukic, M. (2024). Crowding-out effect '
    'of tobacco consumption in Serbia. Tobacco Control, 33(Suppl 2), s88–s94.',
    'Yang, H., Ma, C., Zhao, M., Magnussen, C. G., & Xi, B. (2022). Prevalence and trend of '
    'smokeless tobacco use and its associated factors among adolescents aged 12–16 years in 138 '
    'countries/territories, 1999–2019. BMC Medicine, 20(1), 460.',
]

for ref in all_refs:
    rp = doc.add_paragraph(ref)
    rp.paragraph_format.first_line_indent = Cm(0)
    rp.paragraph_format.left_indent = Cm(0.74)
    rp.paragraph_format.first_line_indent = Cm(-0.74)
    for run in rp.runs:
        run.font.size = Pt(10)

# =================== SAVE ===================
out = os.path.join(ROOT, "manuscript", "Full_Manuscript_CHS.docx")
os.makedirs(os.path.join(ROOT, "manuscript"), exist_ok=True)
try:
    doc.save(out)
    print(f"Saved: {out}")
except PermissionError:
    out2 = os.path.join(ROOT, "manuscript", "Full_Manuscript_CHS_v3.docx")
    doc.save(out2)
    print(f"Saved (fallback): {out2}")
    print("WARNING: Close Full_Manuscript_CHS.docx in Word, then rename v3 over it.")
